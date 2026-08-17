# -*- coding: utf-8 -*-
"""将会计学范文docx转换为md格式，保留表格、章节结构、公式(转LaTeX)与图片(转引用)"""
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn

src = '数字化转型对企业资本结构优化的影响研究——基于A股上市公司的实证检验.docx'
dst = '数字化转型对企业资本结构优化的影响研究——基于A股上市公司的实证检验.md'

doc = Document(src)
body = doc.element.body
paras = doc.paragraphs
tables = doc.tables
pi = 0; ti = 0

# 图片关系：rId -> media文件名（用于判断段落是否含图片及提取文件名）
rels = doc.part.rels
# 按图片出现顺序映射到 figures/ 下可读文件名
_FIG_NAMES = ['fig1_route.png', 'fig2_dt_trend.png', 'fig3_lev_trend.png',
              'fig5_vars_dist.png', 'fig6_corr_heat.png', 'fig7_base_coef.png',
              'fig8_robust_compare.png', 'fig4_mechanism.png',
              'fig10_medi_share.png', 'fig9_hetero_coef.png']
_fig_counter = 0

lines = []

def is_heading(text):
    """识别章节标题"""
    t = text.strip()
    if not t: return 0
    # 一级：第X章
    if t.startswith('第') and '章' in t[:5] and len(t) < 40: return 1
    # 特殊章节
    for k in ['摘要','Abstract','目录','参考文献','附录','致谢']:
        if t.startswith(k) and len(t) < 20: return 1
    # 二级：X.Y 编号 或 一、二、
    if re.match(r'^[一二三四五六七八九十]+[、．.]', t) and len(t) < 50: return 2
    if re.match(r'^\d+\.\d+\s', t) and len(t) < 60: return 2
    if re.match(r'^\d+\.\d+\.\d+\s', t) and len(t) < 60: return 3
    return 0

def has_drawing(p):
    """判断段落是否含图片(w:drawing -> a:blip)"""
    return p._p.findall('.//' + qn('w:drawing'))

def para_to_latex(p):
    """将带真下标/上标的公式段落转为LaTeX行内数学式。
    有下标run时返回 '$...$' 字符串，否则返回None。
    """
    segs = []  # (text, 'sub'|'sup'|'normal')
    has_sub = False
    for r in p.runs:
        txt = r.text
        if txt == '':
            continue
        if r.font.subscript:
            segs.append((txt, 'sub')); has_sub = True
        elif r.font.superscript:
            segs.append((txt, 'sup'))
        else:
            segs.append((txt, 'normal'))
    if not has_sub:
        return None
    # 拼装LaTeX：下标 _{...}，上标 ^{...}，希腊字母转义，×转\cdot
    greek = {'α':'\\alpha','β':'\\beta','γ':'\\gamma','δ':'\\delta',
             'ε':'\\varepsilon','μ':'\\mu','λ':'\\lambda','θ':'\\theta',
             'σ':'\\sigma','ρ':'\\rho','π':'\\pi'}
    # 希腊字母带Unicode下标，整体转 LaTeX 下标形式，如 β₁ -> \beta_1
    greek_sub = {'β₁':'\\beta_1','β₂':'\\beta_2','β₃':'\\beta_3',
                 'α₁':'\\alpha_1','γ₁':'\\gamma_1','δ₁':'\\delta_1'}
    out = []
    for txt, kind in segs:
        # 先处理带下标的希腊字母组合
        for k, v in greek_sub.items():
            txt = txt.replace(k, v)
        # 处理希腊字母
        for k, v in greek.items():
            txt = txt.replace(k, v)
        # 下标数字₁₂₃转普通数字（已由 greek_sub 处理的组合不影响）
        subs = {'₁':'1','₂':'2','₃':'3','₀':'0'}
        for k, v in subs.items():
            txt = txt.replace(k, v)
        # × 转 \cdot
        txt = txt.replace('×', '\\cdot ')
        if kind == 'sub':
            out.append('_{' + txt.strip() + '}')
        elif kind == 'sup':
            out.append('^{' + txt.strip() + '}')
        else:
            out.append(txt)
    latex = ''.join(out).strip()
    # 去除多余空格
    latex = re.sub(r'\s+', ' ', latex)
    return '$' + latex + '$'

def omml_to_latex(p):
    """将 Word 原生 OMML 公式（m:oMath）段落转为 LaTeX 数学式。
    支持 sSub / sSup / sSubSup / r（文本运行）。返回 '$...$' 或 None。
    同时提取段落中非公式部分的编号（如 (4-1)）追加到公式末尾。
    """
    ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    ommls = p._p.findall('.//m:oMath', ns)
    if not ommls:
        return None
    out = []
    for omml in ommls:
        out.append(omml_node_to_latex(omml, ns))
    # 提取公式编号：段落中 oMath 之外的 run 文本
    eq_num = ''
    for r in p.runs:
        t = r.text.strip()
        if t and re.match(r'^\(\d+-\d+\)$', t):
            eq_num = t
            break
    body = ' '.join(out)
    if eq_num:
        body = body + r'\qquad (' + eq_num.strip('()') + ')'
    return '$' + body + '$'


def omml_node_to_latex(el, ns):
    """递归解析单个 OMML 元素为 LaTeX 片段"""
    tag = el.tag.split('}')[-1]
    if tag == 'r':
        # m:r -> m:t 文本
        ts = el.findall('m:t', ns)
        txt = ''.join((t.text or '') for t in ts)
        return omml_escape(txt)
    if tag == 'sSub':
        e = el.find('m:e', ns); sub = el.find('m:sub', ns)
        return omml_node_to_latex(e, ns) + '_{' + omml_node_to_latex(sub, ns) + '}'
    if tag == 'sSup':
        e = el.find('m:e', ns); sup = el.find('m:sup', ns)
        return omml_node_to_latex(e, ns) + '^{' + omml_node_to_latex(sup, ns) + '}'
    if tag == 'sSubSup':
        e = el.find('m:e', ns); sub = el.find('m:sub', ns); sup = el.find('m:sup', ns)
        return (omml_node_to_latex(e, ns) + '_{' + omml_node_to_latex(sub, ns) + '}'
                + '^{' + omml_node_to_latex(sup, ns) + '}')
    # 其他容器（e/sub/sup/num/den 等）：递归子节点
    parts = []
    for child in el:
        parts.append(omml_node_to_latex(child, ns))
    return ''.join(parts)


def omml_escape(txt):
    """将 OMML 文本转义为 LaTeX（希腊字母、数学符号）"""
    greek = {'α': '\\alpha', 'β': '\\beta', 'γ': '\\gamma', 'δ': '\\delta',
             'ε': '\\varepsilon', 'μ': '\\mu', 'λ': '\\lambda', 'θ': '\\theta',
             'σ': '\\sigma', 'ρ': '\\rho', 'π': '\\pi', 'φ': '\\phi',
             'ω': '\\omega', 'η': '\\eta', 'κ': '\\kappa', 'ν': '\\nu',
             'ζ': '\\zeta', 'ψ': '\\psi', 'τ': '\\tau', 'ξ': '\\xi',
             'Α': 'A', 'Β': 'B', 'Γ': '\\Gamma', 'Δ': '\\Delta', 'Σ': '\\Sigma',
             'Λ': '\\Lambda', 'Ω': '\\Omega', 'Π': '\\Pi', 'Φ': '\\Phi'}
    # 先处理希腊字母，避免重复替换
    for k, v in greek.items():
        txt = txt.replace(k, v)
    # 数学符号
    txt = txt.replace('×', '\\times ')
    txt = txt.replace('−', '-')
    txt = txt.replace('·', '\\cdot ')
    txt = txt.replace('∞', '\\infty')
    txt = txt.replace('∈', '\\in ')
    txt = txt.replace('≠', '\\neq ')
    txt = txt.replace('≤', '\\leq ')
    txt = txt.replace('≥', '\\geq ')
    txt = txt.replace('→', '\\rightarrow ')
    txt = txt.replace('≈', '\\approx ')
    # LaTeX 特殊字符转义
    for ch in '#$%&_{}':
        txt = txt.replace(ch, '\\' + ch)
    return txt


def cell_to_md(cell):
    """单元格转md文本：保留下标(_{}形式)与上标(^{}形式)，换行转<br>"""
    segs = []
    for p in cell.paragraphs:
        buf = ''
        for r in p.runs:
            if not r.text:
                continue
            if r.font.subscript:
                buf += '_{' + r.text + '}'
            elif r.font.superscript:
                buf += '^{' + r.text + '}'
            else:
                buf += r.text
        segs.append(buf.strip())
    text = '<br>'.join(s for s in segs if s)
    return text.replace('|', '/')

def table_to_md(tbl):
    """表格转md"""
    rows = []
    for row in tbl.rows:
        cells = [cell_to_md(c) for c in row.cells]
        rows.append(cells)
    if not rows: return ''
    md = []
    ncol = max(len(r) for r in rows)
    for r in rows:
        r = r + ['']*(ncol-len(r))
        md.append('| ' + ' | '.join(r) + ' |')
    sep = '| ' + ' | '.join(['---']*ncol) + ' |'
    md.insert(1, sep)
    return '\n'.join(md)

for child in body.iterchildren():
    if child.tag == qn('w:p'):
        p = paras[pi]; pi += 1
        t = p.text.strip()
        # 优先：图片段落
        if has_drawing(p):
            fname = _FIG_NAMES[_fig_counter] if _fig_counter < len(_FIG_NAMES) else f'image{_fig_counter+1}.png'
            _fig_counter += 1
            lines.append('')
            lines.append(f'![图片](figures/{fname})')
            lines.append('')
            continue
        if not t: continue
        lvl = is_heading(t)
        # 公式段落：优先处理 Word 原生 OMML 公式，其次处理真下标文本公式
        latex = omml_to_latex(p) or para_to_latex(p)
        if latex:
            lines.append('')
            lines.append(f'$$\n{latex.strip("$")}\n$$')
            lines.append('')
            continue
        # 表题/图题
        if t.startswith('表') and (len(t) < 60):
            lines.append(f'\n**{t}**\n')
        elif t.startswith('图') and (len(t) < 60):
            lines.append(f'\n**{t}**\n')
        elif lvl == 1:
            lines.append(f'\n# {t}\n')
        elif lvl == 2:
            lines.append(f'\n## {t}\n')
        elif lvl == 3:
            lines.append(f'\n### {t}\n')
        else:
            lines.append(t + '\n')
    elif child.tag == qn('w:tbl'):
        lines.append('\n' + table_to_md(tables[ti]) + '\n')
        ti += 1

with open(dst, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('已生成:', dst)
print('大小:', round(os.path.getsize(dst)/1024, 1), 'KB')
print('总行数:', len(lines))
print('图片引用数:', _fig_counter)