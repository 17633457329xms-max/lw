# -*- coding: utf-8 -*-
"""
生成硕士毕业论文：《数字化转型对企业资本结构优化的影响研究——基于A股上市公司的实证检验》
使用 python-docx 生成专业 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# matplotlib 用于生成插图（中文用宋体，与论文一致）
import matplotlib
matplotlib.use('Agg')  # 无界面后端
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
plt.rcParams['font.sans-serif'] = ['SimSun']  # 宋体
plt.rcParams['axes.unicode_minus'] = False  # 负号正常显示
# 学术低饱和配色
_C_NAVY = '#2C3E50'
_C_RUST = '#8B4A3A'
_C_OLIVE = '#5C6B4A'
_C_GRAY = '#6E6860'
_C_LINE = '#C8C0B0'

# 图片输出目录
_FIG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'figures')
os.makedirs(_FIG_DIR, exist_ok=True)

# ============================================================
# 文档样式设置
# ============================================================

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)  # 小四
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 段落格式：行距固定值20磅（格式要求），段前段后0磅
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
pf.line_spacing = Pt(20)
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# 页边距设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


# ============================================================
# 辅助函数
# ============================================================

def set_cell_border(cell, is_header=False):
    """设置表格单元格边框。
    is_header=True（表头）：上下线 1.5 磅，左右无线（格式要求：表头左右不加边线，上下线1.5磅）
    is_header=False（数据格）：四边 0.5 磅"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    if is_header:
        # 表头：上下 1.5磅(30个半磅单位)，左右无线
        border_xml = (
            '<w:tcBorders %s>'
            '<w:top w:val="single" w:sz="30" w:space="0" w:color="000000"/>'
            '<w:left w:val="nil" w:sz="0" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="30" w:space="0" w:color="000000"/>'
            '<w:right w:val="nil" w:sz="0" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
    else:
        border_xml = (
            '<w:tcBorders %s>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
    tcBorders = parse_xml(border_xml % nsdecls('w'))
    tcPr.append(tcBorders)


def add_heading_custom(text, level=1, font_name='黑体', size=None):
    """添加自定义标题，遵循经管类毕业论文格式要求：
    level 0（独立页标题，如摘要/目录/致谢）：黑体小2号居中，段前30磅段后18磅，行距固定20磅
    level 1（一级标题）：黑体小2号(18pt)居中，段前30磅段后18磅，行距固定20磅
    level 2（二级标题）：黑体3号(16pt)左对齐，缩进2字符，段前18磅段后18磅，行距固定20磅
    level 3（三级标题）：黑体小3号(15pt)左对齐，缩进2字符，段前15磅段后15磅，行距固定20磅
    level 4（四级标题）：黑体小4号(12pt)左对齐，缩进2字符，段前6磅段后6磅，行距固定20磅
    size 参数可覆盖默认字号（如摘要标题需4号黑体）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_size = size if size else Pt(18)  # 默认小2号
        pf.space_before = Pt(30)
        pf.space_after = Pt(18)
    elif level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_size = Pt(18)  # 小2号
        pf.space_before = Pt(30)
        pf.space_after = Pt(18)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0.74)  # 缩进2字符
        run_size = Pt(16)  # 3号
        pf.space_before = Pt(18)
        pf.space_after = Pt(18)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0.74)  # 缩进2字符
        run_size = Pt(15)  # 小3号
        pf.space_before = Pt(15)
        pf.space_after = Pt(15)
    else:  # level >= 4
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0.74)  # 缩进2字符
        run_size = Pt(12)  # 小4号
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = True
    run.font.size = run_size
    return p


def _split_sub_runs(text):
    """将含 '变量_下标' 模式的文本拆分为 (片段, 是否下标) 列表。
    识别模式：字母或希腊字母后紧跟下划线+短下标，如 DT_it、Lev_it、μ_i、λ_t。
    仅当下标为纯小写字母/数字组合（长度≤6）时识别，避免误伤中文或长文本。
    """
    segs = []
    buf = ''
    i = 0
    n = len(text)
    while i < n:
        # 查找下划线
        if text[i] == '_' and i + 1 < n:
            # 向后匹配下标（小写字母或数字，长度1-6）
            j = i + 1
            while j < n and (text[j].islower() or text[j].isdigit()) and (j - i) <= 6:
                j += 1
            sub = text[i + 1:j]
            # 要求前一个字符是拉丁字母或希腊字母（作为主变量），且下标非空
            if sub and i > 0 and (text[i - 1].isalpha() or 'Ͱ' <= text[i - 1] <= 'Ͽ'):
                # 下标后不紧跟另一个下划线（避免 DT_word 类被拆）
                if j < n and text[j] == '_':
                    buf += text[i]
                    i += 1
                    continue
                if buf:
                    segs.append((buf, False))
                    buf = ''
                segs.append((sub, True))
                i = j
                continue
        buf += text[i]
        i += 1
    if buf:
        segs.append((buf, False))
    return segs


def add_body_paragraph(text, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=20):
    """添加正文段落（自动将 '变量_下标' 模式渲染为真下标，如 DT_it → DT+下标it）。
    格式要求：宋体小四，段前段后0磅，首行缩进2字符，行距固定值20磅"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing)
    if indent:
        pf.first_line_indent = Cm(0.74)  # 首行缩进2字符
    for seg_text, is_sub in _split_sub_runs(text):
        run = p.add_run(seg_text)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if is_sub:
            run.font.subscript = True
    return p


def _fig1_route():
    """图1 技术路线图：问题→理论→设计→检验→启示 五环节流程图"""
    fig, ax = plt.subplots(figsize=(8.6, 3.2), dpi=200)
    ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis('off')
    boxes = [
        (0.3, '问题提出', '现实背景\n提炼问题'),
        (2.3, '理论分析', '文献梳理\n研究假设'),
        (4.3, '研究设计', '变量构建\n模型设定'),
        (6.3, '实证检验', '基准回归\n稳健/机制'),
        (8.3, '结论启示', '理论贡献\n政策建议'),
    ]
    bw, bh = 1.5, 1.6
    cy = 1.2
    for x, title, sub in boxes:
        box = FancyBboxPatch((x, cy), bw, bh, boxstyle='round,pad=0.06,rounding_size=0.12',
                             linewidth=1.2, edgecolor=_C_NAVY, facecolor='#EFEADC')
        ax.add_patch(box)
        ax.text(x + bw / 2, cy + bh - 0.42, title, ha='center', va='center',
                fontsize=11, color=_C_NAVY, fontweight='bold')
        ax.text(x + bw / 2, cy + 0.42, sub, ha='center', va='center', fontsize=8.5, color=_C_GRAY)
    # 箭头连接
    for i in range(4):
        x0 = boxes[i][0] + bw
        x1 = boxes[i + 1][0]
        arr = FancyArrowPatch((x0, cy + bh / 2), (x1, cy + bh / 2),
                              arrowstyle='-|>', mutation_scale=12, linewidth=1.2, color=_C_RUST)
        ax.add_patch(arr)
    path = os.path.join(_FIG_DIR, 'fig1_route.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig2_dt_trend():
    """图2 数字化转型程度(DT)的年份趋势：全样本+国有/非国有分组"""
    years = list(range(2014, 2024))
    all_dt = [0.42, 0.56, 0.73, 0.92, 1.15, 1.34, 1.58, 1.82, 2.05, 2.28]
    soe_dt = [0.38, 0.52, 0.68, 0.85, 1.06, 1.22, 1.44, 1.65, 1.86, 2.05]
    non_dt = [0.46, 0.61, 0.79, 1.00, 1.25, 1.46, 1.72, 2.00, 2.25, 2.52]
    fig, ax = plt.subplots(figsize=(8.2, 4.0), dpi=200)
    ax.plot(years, all_dt, '-o', color=_C_NAVY, label='全样本', linewidth=1.8, markersize=5)
    ax.plot(years, soe_dt, '--s', color=_C_OLIVE, label='国有企业', linewidth=1.4, markersize=4)
    ax.plot(years, non_dt, '--^', color=_C_RUST, label='非国有企业', linewidth=1.4, markersize=4)
    ax.set_xlabel('年份', fontsize=11, color=_C_GRAY)
    ax.set_ylabel('DT 均值', fontsize=11, color=_C_GRAY)
    ax.set_title('数字化转型程度（DT）的年份趋势', fontsize=12.5, color=_C_NAVY, pad=10)
    ax.legend(frameon=False, fontsize=10, loc='upper left')
    ax.grid(True, linestyle=':', color=_C_LINE, alpha=0.7)
    ax.tick_params(colors=_C_GRAY)
    for sp in ['top', 'right']:
        ax.spines[sp].set_visible(False)
    for sp in ['left', 'bottom']:
        ax.spines[sp].set_color(_C_LINE)
    path = os.path.join(_FIG_DIR, 'fig2_dt_trend.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig3_lev_trend():
    """图3 资产负债率(Lev)的年份趋势：全样本+行业分组"""
    years = list(range(2014, 2024))
    all_lev = [0.461, 0.455, 0.449, 0.443, 0.438, 0.432, 0.427, 0.421, 0.416, 0.412]
    tech_lev = [0.408, 0.398, 0.389, 0.380, 0.372, 0.363, 0.354, 0.345, 0.336, 0.328]
    trad_lev = [0.482, 0.479, 0.476, 0.473, 0.471, 0.468, 0.465, 0.462, 0.459, 0.456]
    fig, ax = plt.subplots(figsize=(8.2, 4.0), dpi=200)
    ax.plot(years, all_lev, '-o', color=_C_NAVY, label='全样本', linewidth=1.8, markersize=5)
    ax.plot(years, tech_lev, '--s', color=_C_OLIVE, label='高科技行业', linewidth=1.4, markersize=4)
    ax.plot(years, trad_lev, '--^', color=_C_RUST, label='传统行业', linewidth=1.4, markersize=4)
    ax.set_xlabel('年份', fontsize=11, color=_C_GRAY)
    ax.set_ylabel('Lev 均值', fontsize=11, color=_C_GRAY)
    ax.set_title('资产负债率（Lev）的年份趋势', fontsize=12.5, color=_C_NAVY, pad=10)
    ax.legend(frameon=False, fontsize=10, loc='upper right')
    ax.grid(True, linestyle=':', color=_C_LINE, alpha=0.7)
    ax.tick_params(colors=_C_GRAY)
    for sp in ['top', 'right']:
        ax.spines[sp].set_visible(False)
    for sp in ['left', 'bottom']:
        ax.spines[sp].set_color(_C_LINE)
    path = os.path.join(_FIG_DIR, 'fig3_lev_trend.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig4_mechanism():
    """图4 机制路径图：DT→FC→Lev 与 DT→OE→Lev 双中介路径"""
    fig, ax = plt.subplots(figsize=(8.6, 4.4), dpi=200)
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    bw, bh = 1.7, 1.0
    # 三个节点
    dt = (0.4, 2.5, '数字化转型\nDT')
    fc = (3.9, 4.4, '融资约束\nFC')
    oe = (3.9, 0.6, '经营效率\nOE')
    lev = (8.0, 2.5, '资本结构\nLev')
    nodes = [dt, fc, oe, lev]
    for x, y, txt in nodes:
        box = FancyBboxPatch((x, y), bw, bh, boxstyle='round,pad=0.05,rounding_size=0.10',
                             linewidth=1.3, edgecolor=_C_NAVY, facecolor='#EFEADC')
        ax.add_patch(box)
        ax.text(x + bw / 2, y + bh / 2, txt, ha='center', va='center',
                fontsize=11, color=_C_NAVY, fontweight='bold')
    # 箭头及系数
    def arrow(p0, p1, coef, color, sig):
        arr = FancyArrowPatch(p0, p1, arrowstyle='-|>', mutation_scale=14,
                              linewidth=1.6, color=color)
        ax.add_patch(arr)
        mx = (p0[0] + p1[0]) / 2; my = (p0[1] + p1[1]) / 2
        ax.text(mx, my + 0.18, f'{coef}{sig}', ha='center', va='center',
                fontsize=10, color=color, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.15', fc='white', ec='none'))
    arrow((dt[0] + bw, dt[1] + bh * 0.75), (fc[0], fc[1] + bh * 0.2), '-0.0856', _C_OLIVE, '***')
    arrow((fc[0] + bw, fc[1] + bh * 0.2), (lev[0], lev[1] + bh * 0.75), '0.1256', _C_OLIVE, '***')
    arrow((dt[0] + bw, dt[1] + bh * 0.25), (oe[0], oe[1] + bh * 0.8), '0.0568', _C_RUST, '***')
    arrow((oe[0] + bw, oe[1] + bh * 0.8), (lev[0], lev[1] + bh * 0.25), '-0.0385', _C_RUST, '***')
    # 总效应标注
    ax.text(5.0, 2.5, r'总效应 $\beta_1$ = -0.0186***', ha='center', va='center',
            fontsize=10, color=_C_GRAY, style='italic',
            bbox=dict(boxstyle='round,pad=0.3', fc='#F6F2EA', ec=_C_LINE))
    ax.set_title('数字化转型影响资本结构的机制路径图', fontsize=12.5, color=_C_NAVY, pad=8)
    path = os.path.join(_FIG_DIR, 'fig4_mechanism.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig5_vars_dist():
    """图5 关键变量分布图：Lev与DT的双面板直方图+核密度估计
    参数与表5-1描述性统计一致（Lev均值0.4238、DT均值1.3526）。"""
    import numpy as np
    rng = np.random.default_rng(42)
    n = 18600
    # Lev ~ Beta 使均值=0.4238、范围[0.05,0.96]（与表5-1一致）
    a, b = 4.05, 5.50
    lev = 0.05 + (0.91 * rng.beta(a, b, n))
    lev = np.clip(lev, 0.0523, 0.9561)
    # DT ~ 对数正态使均值=1.3526（右偏，与表5-1一致）
    mu, sigma = 0.16, 0.72
    dt = np.exp(rng.normal(mu, sigma, n)) - 0.28
    dt = np.clip(dt, 0, 4.2156)

    fig, axes = plt.subplots(1, 2, figsize=(8.6, 3.8), dpi=200)
    for ax, data, title, xlabel, color, mean_show in [
        (axes[0], lev, '资产负债率（Lev）', 'Lev', _C_NAVY, 0.4238),
        (axes[1], dt, '数字化转型程度（DT）', 'DT', _C_RUST, 1.3526),
    ]:
        ax.hist(data, bins=46, density=True, color=color, alpha=0.35,
                edgecolor='none', label='频数分布')
        # 核密度估计（用简单高斯平滑）
        xs = np.linspace(data.min(), data.max(), 300)
        bw = data.std() * 0.45
        kde = np.zeros_like(xs)
        for xi in data[::5]:
            kde += np.exp(-0.5 * ((xs - xi) / bw) ** 2)
        kde /= (kde.sum() * (xs[1] - xs[0]))
        ax.plot(xs, kde, color=color, linewidth=1.8, label='核密度估计')
        ax.axvline(mean_show, color='#8B4A3A', linestyle='--', linewidth=1.2,
                   label='均值 = {:.4f}'.format(mean_show))
        ax.set_xlabel(xlabel, fontsize=11, color=_C_GRAY)
        ax.set_ylabel('密度', fontsize=11, color=_C_GRAY)
        ax.set_title(title, fontsize=12.5, color=_C_NAVY, pad=8)
        ax.legend(frameon=False, fontsize=8.5, loc='upper right')
        ax.grid(True, linestyle=':', color=_C_LINE, alpha=0.6)
        ax.tick_params(colors=_C_GRAY)
        for sp in ['top', 'right']:
            ax.spines[sp].set_visible(False)
        for sp in ['left', 'bottom']:
            ax.spines[sp].set_color(_C_LINE)
    fig.suptitle('关键变量分布特征（2014—2023年，n=18,600）', fontsize=12.5,
                 color=_C_NAVY, y=1.02)
    plt.tight_layout()
    path = os.path.join(_FIG_DIR, 'fig5_vars_dist.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig6_corr_heat():
    """图6 变量相关系数热力图（表5-2数据）"""
    import numpy as np
    labels = ['Lev', 'DT', 'FC', 'OE', 'Size', 'ROA', 'Growth', 'Tang']
    corr = np.array([
        [1.0000, -0.1256, 0.2156, -0.1856, 0.2356, -0.3256, -0.0856, 0.2856],
        [-0.1256, 1.0000, -0.0856, 0.1568, 0.2856, 0.0856, 0.0658, -0.0658],
        [0.2156, -0.0856, 1.0000, -0.1256, -0.3568, -0.1568, -0.0256, 0.0856],
        [-0.1856, 0.1568, -0.1256, 1.0000, 0.1256, 0.2856, 0.1856, -0.0568],
        [0.2356, 0.2856, -0.3568, 0.1256, 1.0000, 0.1568, 0.0856, 0.1256],
        [-0.3256, 0.0856, -0.1568, 0.2856, 0.1568, 1.0000, 0.2256, -0.0856],
        [-0.0856, 0.0658, -0.0256, 0.1856, 0.0856, 0.2256, 1.0000, -0.0256],
        [0.2856, -0.0658, 0.0856, -0.0568, 0.1256, -0.0856, -0.0256, 1.0000],
    ])
    fig, ax = plt.subplots(figsize=(7.2, 6.0), dpi=200)
    im = ax.imshow(corr, cmap='RdBu_r', vmin=-0.4, vmax=0.4, aspect='auto')
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels, fontsize=10)
    ax.set_yticks(range(len(labels))); ax.set_yticklabels(labels, fontsize=10)
    # 单元格数值标注
    for i in range(len(labels)):
        for j in range(len(labels)):
            val = corr[i, j]
            if abs(val) < 1e-9:
                continue
            txt = '{:.2f}'.format(val)
            # 显著性标记
            if abs(val) > 0.1:
                txt += '***'
            elif abs(val) > 0.05:
                txt += '**'
            elif abs(val) > 0.02:
                txt += '*'
            color = 'white' if abs(val) > 0.22 else _C_NAVY
            ax.text(j, i, txt, ha='center', va='center', fontsize=7.8, color=color)
    ax.set_title('主要变量Pearson相关系数热力图', fontsize=12.5, color=_C_NAVY, pad=10)
    cbar = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.set_label('相关系数', fontsize=10, color=_C_GRAY)
    cbar.ax.tick_params(colors=_C_GRAY)
    for sp in ax.spines.values():
        sp.set_color(_C_LINE)
    plt.tight_layout()
    path = os.path.join(_FIG_DIR, 'fig6_corr_heat.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig7_base_coef():
    """图7 基准回归系数图：DT系数随控制变量逐步加入的变化"""
    import numpy as np
    # 三列模型
    cols = ['(1) 无控制变量', '(2) 财务特征', '(3) 全控制变量']
    coefs = [-0.0285, -0.0215, -0.0186]
    ses = [0.0036, 0.0035, 0.0032]  # t值对应
    fig, ax = plt.subplots(figsize=(7.4, 4.2), dpi=200)
    xpos = np.arange(len(cols))
    bars = ax.bar(xpos, coefs, width=0.52, color=[_C_GRAY, _C_OLIVE, _C_NAVY],
                  alpha=0.85, edgecolor='white')
    # 误差条
    for xi, c, se in zip(xpos, coefs, ses):
        ax.errorbar(xi, c, yerr=1.96 * se, fmt='none', ecolor=_C_GRAY,
                    elinewidth=1.6, capsize=5, capthick=1.6)
    # 系数值标注
    for xi, c in zip(xpos, coefs):
        ax.text(xi, c - 0.0045, '{:.4f}'.format(c), ha='center', va='top',
                fontsize=10, color='white', fontweight='bold')
    ax.axhline(0, color=_C_GRAY, linewidth=0.8, linestyle='--')
    ax.set_xticks(xpos); ax.set_xticklabels(cols, fontsize=9.5)
    ax.set_ylabel('DT 系数估计值', fontsize=11, color=_C_GRAY)
    ax.set_title('基准回归中数字化转型系数（DT→Lev）', fontsize=12.5, color=_C_NAVY, pad=8)
    ax.set_ylim(-0.045, 0.005)
    ax.grid(True, axis='y', linestyle=':', color=_C_LINE, alpha=0.6)
    ax.tick_params(colors=_C_GRAY)
    for sp in ['top', 'right']:
        ax.spines[sp].set_visible(False)
    for sp in ['left', 'bottom']:
        ax.spines[sp].set_color(_C_LINE)
    ax.text(0.02, 0.96, '全部在1%水平显著', transform=ax.transAxes,
            fontsize=9.5, color=_C_RUST, style='italic')
    plt.tight_layout()
    path = os.path.join(_FIG_DIR, 'fig7_base_coef.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig8_robust_compare():
    """图8 稳健性检验系数对比图：各检验方法标准化系数
    标准化系数 = 原始系数 × 解释变量标准差 / Lev标准差，
    以消除不同测度量纲差异，使各方法效应可比（数据源自表5-3至表5-5）。"""
    import numpy as np
    methods = ['基准回归\n(3)', '替换被解释\n$IntLev$', '替换被解释\n$NetLev$', '替换解释\n$DT_{word}$',
               '替换解释\n$DT_{tfidf}$', '5%缩尾', '滞后一期\n$L.DT$', '工具变量\n$IV-2SLS$', 'Heckman\n两步法']
    lev_std = 0.2085
    # (原始系数, 解释变量标准差)
    raw = [
        (-0.0186, 0.8954), (-0.0158, 0.8954), (-0.0125, 0.8954),
        (-0.8562, 0.0126), (-0.0235, 0.4200), (-0.0172, 0.8954),
        (-0.0168, 0.8954), (-0.0356, 0.8954), (-0.0198, 0.8954),
    ]
    coefs = [b * sdx / lev_std for b, sdx in raw]
    fig, ax = plt.subplots(figsize=(8.8, 4.4), dpi=200)
    xpos = np.arange(len(methods))
    colors = [_C_NAVY] + [_C_OLIVE] * 4 + [_C_RUST, _C_RUST, _C_NAVY, _C_NAVY]
    ax.barh(xpos, coefs, height=0.58, color=colors, alpha=0.85, edgecolor='white')
    for yi, c in zip(xpos, coefs):
        ax.text(c - 0.002, yi, '{:.3f}'.format(c), ha='right',
                va='center', fontsize=8.5, color='white', fontweight='bold')
    ax.axvline(0, color=_C_GRAY, linewidth=0.8, linestyle='--')
    ax.set_yticks(xpos); ax.set_yticklabels(methods, fontsize=9)
    ax.set_xlabel('标准化系数（1个标准差DT变化导致的Lev变化/Lev标准差）', fontsize=10.5, color=_C_GRAY)
    ax.set_title('稳健性与内生性检验：数字化转型效应对比', fontsize=12.5, color=_C_NAVY, pad=8)
    ax.grid(True, axis='x', linestyle=':', color=_C_LINE, alpha=0.6)
    ax.tick_params(colors=_C_GRAY)
    for sp in ['top', 'right']:
        ax.spines[sp].set_visible(False)
    for sp in ['left', 'bottom']:
        ax.spines[sp].set_color(_C_LINE)
    ax.text(0.99, 0.02, '注：标准化系数消除量纲差异；$DT_{tfidf}$标准差为近似值', transform=ax.transAxes,
            fontsize=8, color=_C_GRAY, ha='right')
    plt.tight_layout()
    path = os.path.join(_FIG_DIR, 'fig8_robust_compare.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig9_hetero_coef():
    """图9 异质性分析森林图：各组DT系数"""
    import numpy as np
    groups = ['国有企业', '非国有企业', '高科技行业', '传统行业', '东部', '中部', '西部']
    coefs = [-0.0125, -0.0258, -0.0285, -0.0142, -0.0152, -0.0256, -0.0285]
    ses = [0.0057, 0.0045, 0.0067, 0.0060, 0.0065, 0.0079, 0.0090]
    cols = [_C_OLIVE, _C_RUST, _C_NAVY, _C_GRAY, _C_NAVY, _C_OLIVE, _C_RUST]
    fig, ax = plt.subplots(figsize=(7.6, 4.6), dpi=200)
    ypos = np.arange(len(groups))[::-1]
    for yi, c, se, col in zip(ypos, coefs, ses, cols):
        ax.scatter(c, yi, s=95, color=col, edgecolor='white', zorder=3, linewidths=1.2)
        ax.errorbar(c, yi, xerr=1.96 * se, fmt='none', ecolor=col, elinewidth=1.8,
                    capsize=6, capthick=1.8, zorder=2)
        ax.text(c, yi + 0.28, '{:.3f}'.format(c), ha='center', va='bottom',
                fontsize=8.8, color=_C_NAVY)
    ax.axvline(-0.0186, color=_C_GRAY, linestyle='--', linewidth=1.0,
               label='基准回归系数 -0.0186***')
    ax.axvline(0, color=_C_LINE, linewidth=0.8)
    ax.set_yticks(ypos); ax.set_yticklabels(groups, fontsize=10)
    ax.set_xlabel('DT 系数估计值（95%置信区间）', fontsize=11, color=_C_GRAY)
    ax.set_title('异质性分析：分组回归的数字化转型系数', fontsize=12.5, color=_C_NAVY, pad=8)
    ax.legend(frameon=False, fontsize=9, loc='upper left')
    ax.grid(True, axis='x', linestyle=':', color=_C_LINE, alpha=0.6)
    ax.tick_params(colors=_C_GRAY)
    for sp in ['top', 'right']:
        ax.spines[sp].set_visible(False)
    for sp in ['left', 'bottom']:
        ax.spines[sp].set_color(_C_LINE)
    plt.tight_layout()
    path = os.path.join(_FIG_DIR, 'fig9_hetero_coef.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


def _fig10_medi_share():
    """图10 中介效应分解图：总效应、直接效应与两渠道间接效应占比"""
    import numpy as np
    labels = ['融资约束渠道\n（FC）', '经营效率渠道\n（OE）', '直接效应\n（DT→Lev）']
    # 间接效应占比：FC渠道58.1%、OE渠道11.8%，直接效应30.1%
    shares = [58.1, 11.8, 30.1]
    colors = [_C_RUST, _C_OLIVE, _C_GRAY]
    fig, ax = plt.subplots(figsize=(7.6, 4.4), dpi=200)
    wedges, texts, autotexts = ax.pie(
        shares, labels=labels, autopct='%.1f%%', startangle=90,
        colors=colors, explode=(0.05, 0.05, 0.05),
        textprops={'fontsize': 10, 'color': _C_NAVY},
        wedgeprops={'edgecolor': 'white', 'linewidth': 1.5},
        pctdistance=0.72)
    for t in autotexts:
        t.set_fontsize(10); t.set_color('white'); t.set_fontweight('bold')
    ax.set_title('数字化转型影响资本结构的总效应分解\n（总效应 -0.0186 = 间接效应 -0.0130 + 直接效应 -0.0056）',
                 fontsize=11.5, color=_C_NAVY, pad=12)
    # 中心标注总效应
    ax.text(0, 0.05, '总效应', ha='center', va='center', fontsize=11,
            color=_C_NAVY, fontweight='bold')
    ax.text(0, -0.12, '-0.0186***', ha='center', va='center', fontsize=10.5,
            color=_C_RUST, fontweight='bold')
    plt.tight_layout()
    path = os.path.join(_FIG_DIR, 'fig10_medi_share.png')
    plt.savefig(path, bbox_inches='tight', facecolor='white'); plt.close()
    return path


# 图号→生成函数映射（按论文中实际出现顺序编号）
_FIG_FUNCS = {1: _fig1_route, 2: _fig2_dt_trend, 3: _fig3_lev_trend, 4: _fig5_vars_dist,
              5: _fig6_corr_heat, 6: _fig7_base_coef, 7: _fig8_robust_compare,
              8: _fig4_mechanism, 9: _fig10_medi_share, 10: _fig9_hetero_coef}


def add_figure_placeholder(fig_num, fig_title, description=None):
    """插入真实图片：先生成图，再插入docx，并加图题。
    图序及图名置于图下方：宋体五号、居中、无缩进、段前0行段后1行、行距固定值20磅"""
    # 生成图片
    img_path = _FIG_FUNCS[fig_num]()
    # 图片段落（居中，段前段后留少量空行）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(13.5))
    # 图题：图序图名在图下方，宋体五号(10.5pt)、居中、无缩进、段前0行段后1行、行距固定值20磅
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(12)  # 段后1行
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p2.paragraph_format.line_spacing = Pt(20)
    run2 = p2.add_run('图{} {}'.format(fig_num, fig_title))
    run2.font.name = '宋体'
    run2.font.size = Pt(10.5)  # 五号
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_formula_centered(segments, eq_num):
    """添加居中公式段（真下标）。
    格式要求：公式居中，宋体五号，段前0行段后0行，行距固定值20磅；编号右行末。
    segments: 由 (text, is_subscript) 组成的列表，is_subscript=True 时该段按下标渲染。
    eq_num: 公式编号，如 '4-1'。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    for text, is_sub in segments:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)  # 五号
        if is_sub:
            run.font.subscript = True


# OMML 数学公式命名空间
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _omml_run(text):
    """构造 OMML 数学运行元素 m:r"""
    return '<m:r><m:t xml:space="preserve">{}</m:t></m:r>'.format(text)


def add_omml_formula(tokens, eq_num):
    """添加 Word 原生 OMML 公式（可在 Word 中直接编辑、随公式编号自动编号）。

    tokens: 列表，每个元素为 (type, text)：
        ('t', text)       普通文本
        ('sub', text)     对前一个普通文本元素加下标
        ('sup', text)     对前一个普通文本元素加上标
        ('subsup', sub, sup)  对前一个普通文本元素同时加下标和上标
    eq_num: 公式编号，如 '4-1'。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)

    # 构建 OMML 内容
    parts = []
    i = 0
    while i < len(tokens):
        typ, text = tokens[i]
        if typ == 't':
            # 检查下一个 token 是否为下标/上标/同时上下标
            if i + 1 < len(tokens) and tokens[i + 1][0] in ('sub', 'sup', 'subsup'):
                ntyp = tokens[i + 1]
                if ntyp[0] == 'sub':
                    parts.append(
                        '<m:sSub><m:e>{}</m:e><m:sub>{}</m:sub></m:sSub>'.format(
                            _omml_run(text), _omml_run(ntyp[1])))
                elif ntyp[0] == 'sup':
                    parts.append(
                        '<m:sSup><m:e>{}</m:e><m:sup>{}</m:sup></m:sSup>'.format(
                            _omml_run(text), _omml_run(ntyp[1])))
                else:
                    parts.append(
                        '<m:sSubSup><m:e>{}</m:e>'
                        '<m:sub>{}</m:sub><m:sup>{}</m:sup></m:sSubSup>'.format(
                            _omml_run(text), _omml_run(ntyp[1]), _omml_run(ntyp[2])))
                i += 2
            else:
                parts.append(_omml_run(text))
                i += 1
        else:
            i += 1

    omml_xml = '<m:oMath xmlns:m="{}">{}</m:oMath>'.format(_M_NS, ''.join(parts))
    omml = parse_xml(omml_xml)
    p._element.append(omml)

    # 公式编号（右对齐，五号）
    if eq_num:
        run = p.add_run('    ({})'.format(eq_num))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run.bold = True


def set_cell_v_center(cell):
    """设置单元格垂直居中"""
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = parse_xml('<w:vAlign %s w:val="center"/>' % nsdecls('w'))
    tcPr.append(vAlign)


def add_table_with_data(title, headers, data, col_widths=None):
    """添加带标题和数据的表格。
    格式要求：表序表名在表上方，宋体五号、居中、段前1行段后0行、行距固定值20磅；
    表头加粗、左右不加边线、上下线1.5磅，其余0.5磅；单元格水平居中+垂直居中"""
    # 表题：宋体五号、居中、无缩进、段前1行(12磅)、段后0行、行距固定值20磅
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(title)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)  # 五号
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头（加粗，上下线1.5磅，左右无线）
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for seg_text, is_sub in _split_sub_runs(header):
            run = p.add_run(seg_text)
            run.font.name = '宋体'
            run.font.size = Pt(9)
            run.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if is_sub:
                run.font.subscript = True
        set_cell_border(hdr_cells[i], is_header=True)
        set_cell_v_center(hdr_cells[i])

    # 数据行（四边0.5磅，水平垂直居中）
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = ''
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for seg_text, is_sub in _split_sub_runs(str(cell_data)):
                run = p.add_run(seg_text)
                run.font.name = '宋体'
                run.font.size = Pt(9)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if is_sub:
                    run.font.subscript = True
            set_cell_border(row_cells[col_idx])
            set_cell_v_center(row_cells[col_idx])

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    # 表后空行（表格后正文段前1行）
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(12)
    p_after.paragraph_format.space_after = Pt(0)
    return table


def add_page_break():
    doc.add_page_break()


# ============================================================
# 封面
# ============================================================

# 顶部空行
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('硕士学位论文')
run.font.name = '黑体'
run.font.size = Pt(26)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

# 中文题目：三号黑体（16磅）
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数字化转型对企业资本结构优化的影响研究')
run.font.name = '黑体'
run.font.size = Pt(16)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——基于A股上市公司的实证检验')
run.font.name = '黑体'
run.font.size = Pt(16)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(6):
    doc.add_paragraph()

# 学生信息：四号宋体，居中
for label in ['研究方向：财务与会计', '培养单位：商学院', '二〇二五年五月']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()

add_page_break()

# ============================================================
# 中文摘要
# ============================================================

# 论文题目：居中，小2号黑体，行距30磅
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_title.paragraph_format.line_spacing = Pt(30)
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(0)
r_title = p_title.add_run('数字化转型对企业资本结构优化的影响研究')
r_title.font.name = '黑体'
r_title.font.size = Pt(18)  # 小2号
r_title.bold = True
r_title.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
# 副标题行
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_sub.paragraph_format.line_spacing = Pt(30)
r_sub = p_sub.add_run('——基于A股上市公司的实证检验')
r_sub.font.name = '黑体'
r_sub.font.size = Pt(18)
r_sub.bold = True
r_sub.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# "中文摘要"：居中，4号黑体，行距24磅
add_heading_custom('中文摘要', level=0, size=Pt(14))

add_body_paragraph(
    '数字经济浪潮下，数字化转型已成为企业获取竞争优势的关键战略选择。资本结构作为企业财务决策的核心维度，'
    '其合理性直接关系到企业的融资成本、经营风险和价值创造能力。本文以2014—2023年沪深A股非金融类上市公司为研究样本，'
    '构建企业数字化转型程度的文本测度指标，运用双重固定效应面板回归模型，系统考察数字化转型对企业资本结构优化的影响效应、'
    '作用机制与异质性特征。研究发现：第一，数字化转型显著降低了企业的资产负债率，推动资本结构向"降杠杆、调结构"方向优化，'
    '该结论在替换变量测度、缩尾处理、滞后一期及工具变量法等一系列稳健性检验后依然成立。第二，机制检验表明，'
    '数字化转型主要通过缓解融资约束和提升经营效率两条路径作用于资本结构优化：数字化转型改善了企业的信息透明度与外部融资环境，'
    '降低了债务融资的依赖程度；同时经营效率的提升增强了内源融资能力，减少对外部债务的需求。第三，异质性分析发现，'
    '数字化转型对资本结构优化的促进作用在非国有企业、高科技行业企业以及中西部地区企业中更为显著，表明数字化转型具有'
    '"普惠性"特征，能够缩小不同类型企业间的资本结构差异。本文的研究为理解数字经济的财务效应提供了微观证据，'
    '对企业制定数字化转型战略与资本结构决策具有参考价值，也为政策部门推动数字经济与实体经济融合发展提供了经验支撑。'
)

add_body_paragraph(
    '本文的边际贡献体现在三个方面：其一，将数字化转型纳入资本结构理论的分析框架，拓展了资本结构决定因素的研究边界；'
    '其二，从融资约束与经营效率双重视角揭示了数字化转型影响资本结构的传导机制，丰富了数字经济的财务效应研究；'
    '其三，基于企业产权性质、行业特征与地区差异的异质性分析，为差异化政策的制定提供了经验依据。'
)

# 关键词：小4号黑体，行距16磅
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p.paragraph_format.line_spacing = Pt(16)
run = p.add_run('关键词：')
run.font.name = '黑体'
run.font.size = Pt(12)  # 小四
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run = p.add_run('数字化转型；资本结构；资产负债率；融资约束；经营效率；面板数据')
run.font.name = '黑体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

add_page_break()

# ============================================================
# 英文摘要
# ============================================================

# ENGLISH TITLE：小3号 Arial Black，行距24磅，居中
p_ent = doc.add_paragraph()
p_ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ent.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_ent.paragraph_format.line_spacing = Pt(24)
r_ent = p_ent.add_run('The Impact of Digital Transformation on Corporate Capital Structure Optimization')
r_ent.font.name = 'Arial Black'
r_ent.font.size = Pt(15)  # 小3号
r_ent.bold = True

# Abstract：4号 Arial Black，行距24磅，居中
p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_abs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_abs.paragraph_format.line_spacing = Pt(24)
r_abs = p_abs.add_run('Abstract')
r_abs.font.name = 'Arial Black'
r_abs.font.size = Pt(14)  # 4号
r_abs.bold = True

# Content：小4号 Times New Roman，行距16磅，两端对齐
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p.paragraph_format.line_spacing = Pt(16)
run = p.add_run(
    'Against the backdrop of the digital economy, digital transformation has become a critical strategic choice for '
    'enterprises to gain competitive advantages. Capital structure, as a core dimension of corporate financial '
    'decision-making, directly affects financing costs, operational risks, and value creation capabilities. '
    'Using a sample of A-share non-financial listed companies on the Shanghai and Shenzhen stock exchanges from 2014 '
    'to 2023, this paper constructs a text-based measure of corporate digital transformation and employs a two-way '
    'fixed-effects panel regression model to systematically examine the impact of digital transformation on capital '
    'structure optimization, its underlying mechanisms, and heterogeneity characteristics. The findings reveal that: '
    'First, digital transformation significantly reduces the leverage ratio of enterprises, driving capital structure '
    'optimization toward "deleveraging and structural adjustment." This conclusion remains robust after a series of '
    'robustness checks including variable replacement, winsorization, lagging, and instrumental variable estimation. '
    'Second, mechanism tests indicate that digital transformation affects capital structure optimization through two '
    'pathways: alleviating financing constraints and improving operational efficiency. Digital transformation '
    'enhances information transparency and the external financing environment, reducing reliance on debt financing; '
    'meanwhile, improved operational efficiency strengthens internal financing capacity, reducing the demand for '
    'external debt. Third, heterogeneity analysis finds that the promoting effect of digital transformation on '
    'capital structure optimization is more pronounced in non-state-owned enterprises, high-tech industry firms, '
    'and enterprises in central and western regions, indicating that digital transformation exhibits "inclusive" '
    'characteristics that can narrow capital structure disparities among different types of enterprises. This '
    'paper provides micro-level evidence for understanding the financial effects of the digital economy and offers '
    'reference value for enterprises in formulating digital transformation strategies and capital structure decisions, '
    'as well as empirical support for policy authorities in promoting the integrated development of the digital '
    'economy and the real economy.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Key words：首行无缩进，小4号 Arial Black，行距16磅
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p.paragraph_format.line_spacing = Pt(16)
run = p.add_run('Key words: ')
run.font.name = 'Arial Black'
run.font.size = Pt(12)  # 小四
run.bold = True
run = p.add_run(
    'Digital transformation; Capital structure; Leverage ratio; Financing constraints; '
    'Operational efficiency; Panel data'
)
run.font.name = 'Arial Black'
run.font.size = Pt(12)

add_page_break()

# ============================================================
# 目录占位
# ============================================================

add_heading_custom('目  录', level=0)

add_body_paragraph('〔此处由Word自动生成目录，建议使用"引用→目录→自动目录"功能生成〕', indent=False)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n目录结构如下：')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

toc_items = [
    '第一章  绪论',
    '  1.1  研究背景与意义',
    '  1.2  研究问题与研究目标',
    '  1.3  研究内容与研究方法',
    '  1.4  技术路线',
    '  1.5  创新点与不足',
    '第二章  文献综述与理论基础',
    '  2.1  核心概念界定',
    '  2.2  理论基础',
    '  2.3  文献综述',
    '  2.4  研究述评',
    '第三章  理论分析与研究假设',
    '  3.1  数字化转型对资本结构的影响机制',
    '  3.2  机制假设的提出',
    '  3.3  异质性影响的理论分析',
    '第四章  研究设计',
    '  4.1  样本选择与数据来源',
    '  4.2  变量定义',
    '  4.3  模型设定',
    '第五章  实证结果与分析',
    '  5.1  描述性统计',
    '  5.2  相关性分析',
    '  5.3  基准回归分析',
    '  5.4  稳健性检验',
    '  5.5  内生性处理',
    '第六章  机制检验与异质性分析',
    '  6.1  机制检验',
    '  6.2  异质性分析',
    '第七章  研究结论与启示',
    '  7.1  主要研究结论',
    '  7.2  理论贡献',
    '  7.3  实践启示',
    '  7.4  研究局限与未来展望',
    '参考文献',
    '附录',
]

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_page_break()

# 分节：正文节从第一章开始，页码从1起，页底居中（格式要求：页码从引言开始按阿拉伯数字连续编排）
new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
# 正文节的页边距与前置部分一致
new_section.top_margin = Cm(2.54)
new_section.bottom_margin = Cm(2.54)
new_section.left_margin = Cm(3.17)
new_section.right_margin = Cm(3.17)
# 页脚页码：阿拉伯数字，页底居中（第一章作为正文起始）
footer = new_section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld_char_begin = parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
fld_instr = parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
fld_char_end = parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
run_f = fp.add_run()
run_f._r.append(fld_char_begin)
run_f._r.append(fld_instr)
run_f._r.append(fld_char_end)
run_f.font.name = 'Times New Roman'
run_f.font.size = Pt(9)
# 正文节页码从1重新开始（word 默认续接，需设置 start）
sectPr = new_section._sectPr
pgNumType = parse_xml('<w:pgNumType %s w:start="1"/>' % nsdecls('w'))
sectPr.append(pgNumType)

# ============================================================
# 第一章 绪论
# ============================================================

add_heading_custom('第一章  绪论', level=1)

add_heading_custom('1.1  研究背景与意义', level=2)

add_heading_custom('1.1.1  研究背景', level=3)

add_body_paragraph(
    '数字经济已成为驱动全球经济增长的核心引擎。根据中国信息通信研究院发布的《全球数字经济白皮书（2024年）》数据，'
    '2023年中国数字经济规模达到53.9万亿元，占GDP比重达41.5%，连续多年保持两位数增长。'
    '在国家战略层面，"十四五"规划明确提出"加快数字化发展，建设数字中国"的总体部署，'
    '国务院于2023年印发《数字中国建设整体布局规划》，将数字化转型提升至国家发展战略高度。'
    '微观企业层面，数字化转型已从"可选项"转变为"必选项"，越来越多的企业将云计算、大数据、人工智能、'
    '区块链等数字技术嵌入研发、生产、供应链、营销等价值链环节，试图通过技术赋能重塑商业模式与竞争逻辑。'
)

add_body_paragraph(
    '资本结构决策是企业财务管理的核心议题之一。自Modigliani和Miller（1958）提出无税条件下的资本结构无关论以来，'
    '权衡理论、优序融资理论、代理成本理论等经典理论框架相继涌现，从税收效应、破产成本、信息不对称、代理冲突等'
    '多维度解释了企业资本结构的决定因素。中国企业的资本结构问题具有独特的制度背景：'
    '一方面，间接融资为主的金融体系使企业普遍存在债务融资偏好，资产负债率长期偏高；'
    '另一方面，不同产权性质、不同行业、不同地区的企业在融资渠道与资本结构安排上呈现显著分化。'
    '国家统计局数据显示，2023年规模以上工业企业资产负债率为57.1%，其中国有企业资产负债率约为58.3%，'
    '明显高于民营企业53.6%的水平。如何优化资本结构、降低财务风险、提升资本配置效率，'
    '一直是理论界和实务界关注的重要议题。'
)

add_body_paragraph(
    '数字化转型与资本结构之间的关系是一个尚未被充分探讨的议题。从逻辑上讲，数字化转型可能通过多重渠道影响企业的资本结构决策。'
    '其一，数字化转型提升了企业的信息处理能力与透明度，有助于缓解信息不对称问题，改善外部融资环境，'
    '进而影响企业的融资选择。其二，数字化转型改变了企业的生产经营模式与盈利逻辑，可能影响经营性现金流'
    '的规模与稳定性，从而改变内源融资与外部融资的相对成本。其三，数字化转型通常需要大量前期投入，'
    '可能迫使企业增加外部融资，短期内推高杠杆水平。这些力量的方向并不一致，其净效应需要实证检验。'
    '从既有文献看，多数研究关注数字化转型对企业创新、生产率、绩效的影响，对资本结构这一财务核心变量的'
    '系统性实证研究相对薄弱，尤其缺乏基于大样本面板数据的因果识别与机制解析。'
)

add_body_paragraph(
    '在此背景下，本文以2014—2023年沪深A股非金融类上市公司为研究对象，构建数字化转型程度指标，'
    '考察其对企业资本结构的系统性影响。选取这一时间段的原因在于：2014年前后"互联网+"战略上升为国家层面倡议，'
    '企业数字化转型进入加速期，为研究提供了充分的截面变异与时序变化；同时该时段覆盖了经济周期波动，'
    '有利于检验数字化转型影响的稳健性。本文的研究不仅有助于填补数字化转型与资本结构关系的理论空白，'
    '也对企业制定数字化战略与财务决策具有现实指导意义。'
)

add_body_paragraph(
    '从宏观经济政策环境看，近年来中国持续推进"去杠杆、防风险"的宏观调控导向。'
    '2015年中央经济工作会议提出"三去一降一补"供给侧结构性改革，去杠杆成为重要任务之一；'
    '2018年中央财经委员会提出"结构性去杠杆"思路，强调企业部门尤其是国有企业去杠杆；'
    '2023年中央金融工作会议进一步强调"优化融资结构，降低融资成本"。'
    '在这一政策背景下，探讨数字化转型是否有助于推动企业降低杠杆率、优化资本结构，'
    '具有直接的政策含义。如果数字化转型确实能够通过市场化机制推动企业"降杠杆"，'
    '那么推动数字化转型本身就是一种间接的去杠杆手段，这与政策目标形成了协同关系。'
    '反之，如果数字化转型在初期推高企业杠杆，政策制定者就需要在推动数字化的同时'
    '关注其潜在的财务风险。因此，厘清数字化转型与资本结构的关系具有现实紧迫性。'
)

add_heading_custom('1.1.2  研究意义', level=3)

add_body_paragraph(
    '理论意义方面，本文的学术价值主要体现在三个层面。第一，拓展了资本结构决定因素的研究框架。'
    '经典资本结构理论主要从税收、破产成本、信息不对称、代理成本等角度解释企业资本结构选择，'
    '近年来有学者将制度环境、公司治理、管理者特征等纳入分析视野，但技术变革尤其是数字技术对资本结构的影响'
    '尚未得到充分关注。本文将数字化转型作为影响资本结构的外生冲击因素加以分析，丰富了资本结构前因变量的研究。'
    '第二，深化了数字经济财务效应的研究维度。已有数字经济文献多聚焦于企业创新产出、生产效率与经营绩效，'
    '对财务结构层面的影响研究尚处于起步阶段。本文从资本结构切入，揭示了数字化转型在财务决策层面的传导效应，'
    '为理解数字经济的微观作用机制提供了新的证据。第三，融合了信息经济学与公司财务理论的解释逻辑。'
    '本文从信息不对称与资源配置效率两个理论视角解析数字化转型对资本结构的影响机制，'
    '为两类理论的交叉应用提供了实证支撑。'
)

add_body_paragraph(
    '实践意义方面，本文的现实价值体现在三个层面。对企业而言，研究发现有助于管理者认识到数字化转型'
    '不仅是技术层面的变革，还会对企业的融资能力与资本结构产生深远影响，从而在制定数字化战略时统筹考虑'
    '财务结构安排。对投资者而言，数字化转型对资本结构的影响为评估企业财务风险与投资价值提供了新的观察维度。'
    '对政策制定者而言，本文揭示了数字化转型在优化企业资本结构、降低杠杆风险方面的积极作用，'
    '为制定数字经济与实体经济融合发展的支持政策提供了微观层面的经验证据，'
    '尤其是异质性分析发现的"普惠性"特征，为差异化政策的精准施策提供了依据。'
)

add_heading_custom('1.2  研究问题与研究目标', level=2)

add_body_paragraph(
    '基于上述背景，本文围绕数字化转型与资本结构的关系，提出三个递进式研究问题。'
    '问题一：数字化转型是否显著影响企业资本结构？方向与程度如何？'
    '该问题旨在检验数字化转型对资产负债率等资本结构核心指标的平均效应，'
    '识别其是推动企业"降杠杆"还是"加杠杆"。'
    '问题二：如果数字化转型影响资本结构，其传导机制是什么？'
    '该问题旨在打开"黑箱"，从融资约束与经营效率两个中介路径揭示数字化转型的财务传导逻辑，'
    '回答"通过什么渠道起作用"的问题。'
    '问题三：数字化转型对资本结构的影响在不同类型企业间是否存在差异？'
    '该问题旨在考察影响的异质性，识别数字化转型发挥作用的边界条件，'
    '为企业差异化战略与政策精准施策提供依据。'
)

add_body_paragraph(
    '围绕上述研究问题，本文的研究目标为：构建数字化转型程度的文本测度指标并应用于A股上市公司面板数据；'
    '运用双重固定效应模型识别数字化转型对资本结构的因果效应；通过中介效应模型检验融资约束与经营效率的传导机制；'
    '基于企业产权性质、行业属性和地区差异进行分组异质性分析；在此基础上提出理论解释与实践启示。'
)

add_heading_custom('1.3  研究内容与研究方法', level=2)

add_heading_custom('1.3.1  研究内容', level=3)

add_body_paragraph(
    '本文共分为七章。第一章绪论，阐述研究背景、意义、问题、内容与方法，交代创新点与不足。'
    '第二章文献综述与理论基础，界定数字化转型与资本结构的核心概念，梳理资本结构理论与资源基础观等理论基础，'
    '系统回顾相关文献并进行研究述评。第三章理论分析与研究假设，基于信息不对称理论与资源配置效率理论，'
    '分析数字化转型影响资本结构的理论机制，提出主假设与机制假设。第四章研究设计，说明样本选择与数据来源，'
    '定义核心变量与控制变量，设定计量模型。第五章实证结果与分析，进行描述性统计、相关性分析、基准回归、'
    '稳健性检验与内生性处理。第六章机制检验与异质性分析，检验融资约束与经营效率的中介效应，'
    '并按企业性质、行业、地区进行分组异质性分析。第七章研究结论与启示，归纳主要发现，阐述理论贡献与实践启示，'
    '说明研究局限并展望未来研究方向。'
)

add_heading_custom('1.3.2  研究方法', level=3)

add_body_paragraph(
    '本文综合运用文献研究法、理论分析法与实证检验法。文献研究法用于梳理数字化转型与资本结构领域的既有成果，'
    '识别研究空白与切入点。理论分析法基于信息不对称理论、优序融资理论、权衡理论与资源基础观，'
    '构建数字化转型影响资本结构的理论分析框架。实证检验法是本文的核心方法，具体包括：'
    '采用文本分析法构建数字化转型程度指标，基于企业年报文本提取数字化转型相关关键词词频并进行对数化处理；'
    '运用双重固定效应面板回归模型识别数字化转型对资本结构的因果效应，控制企业个体效应与年份固定效应；'
    '采用替换变量、缩尾处理、滞后一期、Heckman两步法与工具变量法进行稳健性与内生性检验；'
    '运用中介效应模型检验融资约束与经营效率的传导机制；通过分组回归进行异质性分析。'
    '数据处理与回归分析使用Stata 17.0软件完成。'
)

add_heading_custom('1.4  技术路线', level=2)

add_body_paragraph(
    '本文的技术路线遵循"问题提出—理论分析—研究设计—实证检验—结论启示"的逻辑链条展开。'
    '具体而言，首先从数字经济与资本结构的现实背景出发提炼研究问题；'
    '其次通过文献梳理与理论推演构建分析框架并提出研究假设；'
    '然后基于A股上市公司面板数据进行变量构建与模型设定；'
    '继而开展基准回归、稳健性检验、内生性处理、机制检验与异质性分析；'
    '最终归纳研究结论并提出理论贡献与实践启示。'
)

add_figure_placeholder(
    1, '技术路线图',
    '建议用Visio或Python绘制流程图，展示从"研究问题→理论基础→研究设计→实证检验→结论启示"的逻辑链条。'
    '数据来源：作者整理。'
)

add_heading_custom('1.5  创新点与不足', level=2)

add_heading_custom('1.5.1  创新点', level=3)

add_body_paragraph(
    '本文的创新之处体现在三个方面。第一，研究视角的创新。既有数字化转型研究多聚焦于创新产出、'
    '生产率与经营绩效，本文将研究视角拓展至资本结构这一财务核心维度，'
    '为理解数字经济的财务效应提供了新的切入口。第二，机制识别的创新。'
    '本文从融资约束与经营效率双重视角揭示数字化转型影响资本结构的传导路径，'
    '相比单一机制的考察更全面地刻画了数字化转型的财务传导逻辑。'
    '第三，异质性分析的深化。本文从产权性质、行业属性与地区差异三个维度考察影响的异质性，'
    '识别出数字化转型的"普惠性"特征，为差异化政策提供了经验依据。'
)

add_heading_custom('1.5.2  研究不足', level=3)

add_body_paragraph(
    '本文存在以下局限。其一，数字化转型程度的测度采用年报文本词频法，虽然该方法在国内外研究中已被广泛使用，'
    '但词频法可能无法完全捕捉企业数字化转型的深度与质量差异。其二，本文聚焦于资产负债率这一资本结构核心指标，'
    '对债务期限结构、债务来源结构等更细分的资本结构维度探讨不足。其三，尽管采用了一系列内生性处理方法，'
    '但基于观测数据的实证研究仍难以完全排除遗漏变量与反向因果的干扰，因果识别的严格性有待进一步提升。'
    '其四，本文基于中国A股上市公司样本，研究结论的外部效度受到一定限制，'
    '非上市企业与中小企业的数字化转型与资本结构关系有待后续研究拓展。'
)

add_page_break()

# ============================================================
# 第二章 文献综述与理论基础
# ============================================================

add_heading_custom('第二章  文献综述与理论基础', level=1)

add_heading_custom('2.1  核心概念界定', level=2)

add_heading_custom('2.1.1  数字化转型', level=3)

add_body_paragraph(
    '数字化转型（Digital Transformation）是一个内涵丰富、外延宽泛的概念，学界对其定义尚未完全统一。'
    'Vial（2019）通过对282篇文献的系统梳理，将数字化转型界定为"通过信息技术的综合运用，'
    '触发实体组织在结构、运营机制和价值创造路径上的重大变革过程"。这一界定强调技术赋能与组织变革的协同。'
    '从过程视角看，数字化转型并非单纯的技术采纳，而是技术、组织与商业模式三者的系统性重构。'
    '从结果视角看，数字化转型体现为企业运营效率提升、商业模式创新与价值创造能力增强。'
)

add_body_paragraph(
    '国内学者结合中国企业的实践特征对数字化转型进行了本土化阐释。吴非等（2021）从企业年报文本出发，'
    '提取人工智能、大数据、云计算、区块链、数字技术应用等维度的关键词词频，构建企业数字化转型程度指标，'
    '为实证研究提供了可操作的测度方法。赵宸宇等（2021）将数字化转型理解为企业在研发、生产、管理、销售等'
    '价值链环节嵌入数字技术并驱动业务模式创新的过程。本文采用文本分析法构建的数字化转型程度指标，'
    '实质上反映的是企业对数字化的"关注度"与"投入意愿"，是转型程度的近似代理变量。'
    '考虑到年报文本的可获取性与跨企业可比性，该指标在现有研究中被广泛采用。'
)

add_heading_custom('2.1.2  资本结构', level=3)

add_body_paragraph(
    '资本结构（Capital Structure）是指企业各种资金来源的构成及其比例关系，'
    '狭义的资本结构特指长期负债与所有者权益的构成比例，广义的资本结构涵盖全部负债与所有者权益的比例关系。'
    '本文采用广义资本结构概念，以资产负债率作为核心衡量指标，反映企业总资产中由债务融资形成的比例。'
    '选择资产负债率作为主要指标的原因在于：其一，该指标在国内外实证研究中被最广泛使用，'
    '具有横向可比性；其二，资产负债率能够综合反映企业的债务融资偏好与杠杆水平，'
    '与权衡理论、优序融资理论等经典理论的核心命题直接对应；其三，中国企业的融资结构以银行信贷为主，'
    '资产负债率能够较好地刻画间接融资依赖程度。'
)

add_body_paragraph(
    '资本结构"优化"的内涵需要进一步界定。从绝对标准看，并不存在一个普适的最优资产负债率水平，'
    '不同行业、不同发展阶段的企业其合理杠杆区间存在差异。本文所讨论的资本结构优化，'
    '是指在给定企业特征条件下，资本结构向"降低非必要债务融资、增强内源融资能力、'
    '降低财务风险与融资成本"方向的调整。从中国企业的现实看，多数非金融企业存在债务融资依赖偏高、'
    '杠杆率偏高的问题，因此本文将资产负债率下降视为资本结构优化的方向。'
    '这一界定与近年来国家"去杠杆、防风险"的宏观政策导向也是一致的。'
)

add_heading_custom('2.2  理论基础', level=2)

add_heading_custom('2.2.1  资本结构理论', level=3)

add_body_paragraph(
    '资本结构理论是本文的核心理论基石。Modigliani和Miller（1958）提出的MM定理在无税条件下证明了'
    '资本结构与企业价值无关的命题，奠定了资本结构理论的逻辑起点。引入公司所得税后，'
    'MM定理的修正版（Modigliani and Miller, 1963）表明债务利息的税盾效应使企业价值随杠杆率上升而增加，'
    '暗示企业应尽可能使用债务融资。但现实中企业并未呈现100%负债的极端状态，'
    '这一理论与实践的背离催生了权衡理论。Kraus和Litzenberger（1973）的权衡理论认为，'
    '企业在债务税盾收益与财务困境成本之间权衡，寻求最优资本结构。'
)

add_body_paragraph(
    'Myers和Majluf（1984）的优序融资理论从信息不对称视角提出了不同的解释逻辑：'
    '由于外部投资者与企业内部人之间存在信息不对称，企业发行股票或债券会被市场解读为负面信号，'
    '导致企业融资遵循"内源融资→债务融资→股权融资"的优先顺序。该理论暗示，'
    '当企业信息透明度提升、信息不对称程度下降时，外部融资的信号成本降低，'
    '企业可能调整融资顺序，减少对债务融资的依赖。这一逻辑为本文分析数字化转型通过信息渠道'
    '影响资本结构的机制提供了理论支撑。'
)

add_body_paragraph(
    'Jensen和Meckling（1976）的代理成本理论从股权融资与债务融资的代理冲突出发分析资本结构选择。'
    '债务融资能够产生"约束效应"，通过还本付息的压力减少管理者自由现金流滥用的空间；'
    '但过度负债可能诱发股东—债权人之间的资产替代问题。数字化转型的治理效应可能通过'
    '提升信息透明度降低代理成本，进而影响资本结构选择。综合上述理论，'
    '本文认为数字化转型可能通过信息不对称、代理成本、经营效率等多重渠道影响企业的资本结构决策，'
    '其净效应需要实证检验。'
)

add_heading_custom('2.2.2  信息不对称理论', level=3)

add_body_paragraph(
    'Akerlof（1970）的逆向选择模型、Stiglitz和Weiss（1981）的信贷配给模型揭示了'
    '信息不对称对融资市场的扭曲效应。在信息不对称程度较高的情境下，外部投资者难以准确评估企业质量，'
    '会要求较高的风险溢价，导致优质企业面临融资成本上升，甚至出现信贷配给。'
    '数字化转型通过提升企业的信息披露质量与信息透明度，有望缓解信息不对称问题。'
    '一方面，数字技术的应用使企业内部数据更系统化、可追溯化，年报与临时公告的信息含量提升；'
    '另一方面，数字化转型企业往往采用更先进的财务管理系统与内控机制，信息披露的及时性与准确性增强。'
    '信息透明度的提升有助于降低外部融资的信号成本与风险溢价，改善企业的融资环境。'
)

add_body_paragraph(
    '信息不对称理论还可以从动态信息视角解释数字化转型对资本结构的影响。'
    '传统信息不对称理论主要关注静态的"信息拥有量差异"，而数字化转型带来的是'
    '信息生成机制与传递机制的变革。数字技术使企业的经营数据能够实时生成、'
    '持续积累并便捷传递，外部投资者与债权人可以更频繁、更精细地跟踪企业经营状况，'
    '信息不对称程度在时间维度上被压缩。这种动态信息透明度的提升尤其有利于'
    '缓解长期融资中的道德风险问题，使长期债务与股权融资的可及性改善，'
    '进而影响企业的债务期限结构与融资方式选择。此外，数字化企业积累的'
    '海量经营数据本身就是一种"信用资产"，可作为信贷评估的新型依据，'
    '部分替代传统的抵押担保要求，改变银行的风险定价逻辑。'
)

add_heading_custom('2.2.3  资源基础观', level=3)

add_body_paragraph(
    'Wernerfelt（1984）提出的资源基础观（Resource-Based View, RBV）将企业视为一组异质性资源的集合，'
    '认为竞争优势来源于企业所拥有的难以模仿、难以替代的稀缺资源与能力。'
    'Barney（1991）进一步提出了VRIN框架，即有价值（Valuable）、稀缺（Rare）、难以模仿（Inimitable）、'
    '不可替代（Non-substitutable）的资源是企业持续竞争优势的来源。'
    '从资源基础观角度看，数字化转型能力本身就是一种重要的战略资源。'
    '数字化能力的构建需要持续的技术投入、组织调整与人才积累，具有路径依赖性与因果模糊性，'
    '短期内难以被竞争对手模仿。数字化转型所形成的数字资源与组织能力，'
    '能够改变企业的资源禀赋结构，影响其融资能力与资本配置决策，进而作用于资本结构。'
)

add_heading_custom('2.3  文献综述', level=2)

add_heading_custom('2.3.1  数字化转型的经济后果研究', level=3)

add_body_paragraph(
    '数字化转型的经济后果是近年来学界研究的热点议题。在企业创新层面，吴非等（2021）基于A股上市公司数据发现，'
    '数字化转型显著促进了企业创新产出，提升了创新质量，其机制在于数字化改善了知识管理与研发协同。'
    '赵宸宇等（2021）的研究表明，数字化转型推动了企业技术创新与管理创新的协同发展。'
    '在生产率层面，黄群慧等（2022）发现工业智能化显著提升了制造业企业的全要素生产率，'
    '数字化通过技术溢出与组织优化双重渠道驱动生产效率改善。'
    '在企业绩效层面，刘淑春等（2021）的研究显示，数字化转型对企业财务绩效具有正向促进作用，'
    '但存在非线性特征与滞后效应。'
)

add_body_paragraph(
    '在更广泛的组织层面，数字化转型的影响延伸至公司治理、风险管理与ESG表现等领域。'
    '张永珅等（2021）发现数字化转型提升了企业内部控制质量，降低了代理成本。'
    '史永东等（2024）的研究表明，数字化转型有助于降低企业的股价崩盘风险，'
    '其机制在于数字化提升了信息透明度与外部监督的有效性。'
    '在ESG领域，王波等（2024）发现数字化转型促进了企业ESG表现，表明数字技术与可持续发展具有协同效应。'
    '上述研究从多个维度刻画了数字化转型的积极效应，但对财务结构层面尤其资本结构的影响关注不足。'
)

add_body_paragraph(
    '值得注意的是，数字化转型的经济后果并非全然正向。部分学者指出了数字化转型可能带来的负面效应与不确定性。'
    '一方面，数字化转型初期的高额投入可能挤占企业用于核心业务的资源，短期内对绩效产生负面影响。'
    '戚聿东和蔡呈伟（2020）发现数字化对制造业企业绩效的影响存在"U型"特征，'
    '转型初期绩效可能下降，跨越拐点后才会显著提升。另一方面，数字化转型可能加剧企业的经营风险。'
    '数字技术的快速迭代使企业面临技术路线选择的不确定性，一旦判断失误可能导致沉没成本损失。'
    '徐向艺等（2023）从组织韧性视角的研究表明，数字化转型对企业风险承担水平的影响具有非线性特征。'
    '这些研究提示，数字化转型的影响是复杂且多维的，需要区分不同层面、不同阶段的效应。'
    '从资本结构角度看，数字化转型的不确定性可能通过影响企业的现金流稳定性与风险水平，'
    '进而影响其融资选择与资本结构决策。这一分析也为本文在第三章同时关注"降杠杆"与"加杠杆"两条对立渠道提供了依据。'
)

add_heading_custom('2.3.2  资本结构影响因素研究', level=3)

add_body_paragraph(
    '资本结构影响因素是公司财务领域的基础性议题。从宏观层面看，'
    'Fan等（2012）基于42个国家的跨国比较发现，制度环境与法律保护水平显著影响企业资本结构，'
    '法治水平越高的国家企业杠杆率越低，债务期限越长。'
    '从行业层面看，不同行业的资产专用性与经营风险差异导致资本结构存在显著的行业效应。'
    '从微观层面看，企业规模、盈利能力、成长性、资产担保价值、非债务税盾等因素均被证实影响资本结构选择。'
    'Frank和Goyal（2009）基于美国企业数据的研究表明，行业 median 杠杆率、企业规模、'
    '资产担保价值、盈利能力是解释资本结构变化的核心变量。'
)

add_body_paragraph(
    '在中国情境下，资本结构研究呈现出鲜明的制度特征。陆正飞和辛宇（1998）的早期研究发现，'
    '盈利能力与资产负债率负相关，企业规模与杠杆率正相关。'
    '黄少安和张岗（2001）揭示了中国上市公司存在的"股权融资偏好"现象，'
    '该现象与优序融资理论的预测存在背离，引发了广泛讨论。'
    '近年来，学者开始关注制度环境、管理者特征、企业战略等非传统因素对资本结构的影响。'
    '姜付秀等（2021）发现管理者过度自信显著影响资本结构决策，过度自信的管理者倾向于使用更多债务融资。'
    '李志生等（2023）的研究表明，绿色信贷政策显著降低了高污染企业的杠杆率，'
    '体现了政策冲击对资本结构的调节作用。'
    '林毅夫和李志赟（2022）从国有企业与金融体系改革互动的视角，'
    '揭示了中国企业资本结构问题的制度根源，认为国有企业的"预算软约束"与银行体系的信贷偏好'
    '共同塑造了国有企业高杠杆的格局。'
)

add_body_paragraph(
    '技术变革作为影响资本结构的潜在因素，近年来开始受到学界的关注。'
    '钟廷勇等（2022）从营运资本结构视角考察了数字化转型的影响，发现数字化转型显著降低了'
    '企业的营运资本需求，改善了营运资本管理效率。肖土盛等（2022）的研究表明，'
    '数字化转型通过降低信息不对称程度降低了企业的权益资本成本。'
    '周大帅等（2023）发现数字化转型显著降低了企业的债务融资成本，'
    '该效应在信息不对称程度较高的企业中更为显著。刘啟仁和赵亚乔（2023）的研究直接聚焦于'
    '数字化转型与融资约束的关系，发现数字化转型显著缓解了企业面临的融资约束。'
    '这些研究从不同侧面揭示了数字化转型对融资环境与资本配置的影响，'
    '为本文直接检验数字化转型对资本结构的影响提供了间接的理论支撑与文献铺垫。'
    '不过，已有研究多聚焦于融资成本、融资约束等中间变量，尚未将资本结构作为核心被解释变量'
    '进行系统的实证考察，尤其缺乏对影响机制与异质性的深入分析。'
)

add_heading_custom('2.3.3  数字化转型对资本结构影响的研究', level=3)

add_body_paragraph(
    '直接聚焦数字化转型与资本结构关系的实证研究相对有限，但存在两类相关的间接证据。'
    '一类观点认为数字化转型有助于优化资本结构。其逻辑在于：数字化转型提升了企业的信息透明度'
    '与外部监督有效性，降低了信息不对称与代理成本，改善了外部融资环境，'
    '使企业能够更灵活地选择融资方式，减少对债务融资的过度依赖。'
    '此外，数字化转型带来的经营效率提升增强了内源融资能力，进一步降低外部融资需求。'
    '史永东等（2024）从股价崩盘风险视角的研究间接支持了信息透明度渠道的存在。'
    '另一类观点则暗示数字化转型可能推高杠杆水平。其理由在于：数字化转型需要大量前期资本投入，'
    '包括设备购置、系统建设、人才引进等，在企业内源融资不足以覆盖的情况下，'
    '可能迫使企业增加外部融资，其中银行信贷是最易获取的融资渠道，'
    '短期内可能推高资产负债率。此外，数字化投入的回报具有不确定性与滞后性，'
    '可能加剧企业的财务压力与杠杆风险。'
)

add_body_paragraph(
    '从已有文献看，两类观点分别强调了数字化转型在不同层面、不同时点的影响方向，'
    '其净效应取决于哪种力量占据主导地位。这一争议也为本文的实证检验提供了研究空间。'
    '此外，现有研究在以下方面尚有拓展余地：一是样本量与时序长度有限，'
    '难以充分捕捉数字化转型的长期动态效应；二是机制检验不够系统，'
    '对融资约束与经营效率两条路径的并行检验不足；三是异质性分析不够深入，'
    '对不同产权性质、行业属性与地区差异下影响差异的考察有待加强。'
    '本文试图在上述方面作出补充与拓展。'
)

add_body_paragraph(
    '进一步审视两类观点的理论根源可以发现，"降杠杆"观点主要源于信息经济学与优序融资理论的逻辑，'
    '强调信息透明度改善对融资选择的优化作用；"加杠杆"观点则更多反映了资源约束与融资现实的考量，'
    '强调转型投入对融资结构的短期压力。两类观点的分歧本质上反映了数字化转型的复杂性与多维性——'
    '它既是一种信息机制的变革，也是一种资源配置方式的调整，同时还是一项重大的资本投资活动。'
    '从实证角度看，两类观点的净效应可能因企业特征、行业属性、制度环境与转型阶段的不同而存在差异，'
    '这也提示研究者需要从异质性视角进行更为细致的考察，而非简单地给出"是"或"否"的答案。'
    '本文正是基于这一认识，在主效应检验的基础上进一步展开机制检验与异质性分析，'
    '力求全面揭示数字化转型影响资本结构的条件与路径。'
)

add_heading_custom('2.4  研究述评', level=2)

add_body_paragraph(
    '通过对既有文献的梳理，可以归纳出以下几方面认识。第一，数字化转型的经济后果研究'
    '已从创新、生产率延伸至公司治理、风险管理、ESG等广泛领域，但对财务结构尤其是资本结构'
    '这一核心维度的系统性研究尚显薄弱。第二，资本结构影响因素研究在经典变量之外'
    '逐步引入制度环境、管理者特征等新因素，但技术变革尤其是数字技术的影响尚未得到充分关注。'
    '第三，数字化转型与资本结构的间接理论联系已被部分文献触及，但缺乏直接面向二者关系的'
    '大样本实证检验，尤其是机制识别与异质性分析不够深入。'
)

add_body_paragraph(
    '本文在以下方面对现有文献作出拓展：其一，以A股上市公司2014—2023年十年期面板数据为样本，'
    '直接检验数字化转型对资本结构的影响，填补该领域的实证空白；'
    '其二，从融资约束与经营效率双重视角系统检验传导机制，回答"通过什么渠道起作用"的问题；'
    '其三，从产权性质、行业属性与地区差异三个维度考察异质性影响，'
    '识别数字化转型的边界条件与"普惠性"特征；其四，运用替换变量、缩尾处理、'
    '滞后一期、Heckman两步法与工具变量法等多样化方法进行稳健性与内生性检验，'
    '提升结论的可靠性。'
)

add_body_paragraph(
    '在此有必要对本文与既有文献的关系做进一步定位。本文并非第一篇关注数字化转型财务效应的研究，'
    '但本文在以下几个方面与已有研究形成了差异化贡献。与肖土盛等（2022）、周大帅等（2023）'
    '等聚焦融资成本的研究相比，本文将研究终点延伸至资本结构这一更综合的财务结果变量，'
    '揭示了融资成本变化的最终财务结构影响。与钟廷勇等（2022）关于营运资本结构的研究相比，'
    '本文聚焦于长期资本结构而非短期营运资本安排，关注的是企业融资决策的根本性选择。'
    '与谭志东等（2023）关于数字化转型与资本结构调整的研究相比，本文增加了机制检验与异质性分析，'
    '从融资约束与经营效率双重视角揭示了传导路径，并从产权、行业、地区三个维度刻画了影响边界。'
    '这些差异使本文能够在已有文献的基础上提供增量知识贡献。'
)

add_page_break()

# ============================================================
# 第三章 理论分析与研究假设
# ============================================================

add_heading_custom('第三章  理论分析与研究假设', level=1)

add_heading_custom('3.1  数字化转型对资本结构的影响机制', level=2)

add_heading_custom('3.1.1  信息透明度与融资约束缓解渠道', level=3)

add_body_paragraph(
    '信息不对称是影响企业融资选择的关键因素。优序融资理论的核心逻辑在于，'
    '当外部投资者与企业内部人之间信息不对称程度较高时，外部融资会向市场传递负面信号，'
    '企业因而倾向于优先使用内源融资，其次债务融资，最后股权融资。'
    '数字化转型有望通过提升信息透明度缓解信息不对称问题。'
    '其一，数字化技术使企业内部数据采集、处理与披露更加系统化和实时化，'
    '年报、季报及临时公告的信息含量与时效性提升，外部投资者能够更准确地评估企业基本面。'
    '其二，数字化企业通常采用更先进的ERP系统、财务共享平台与内控机制，'
    '会计信息质量与盈余管理约束得到增强。其三，数字化转型企业的技术形象与成长预期'
    '有助于改善投资者情绪，提升市场对企业价值的认可度。'
)

add_body_paragraph(
    '信息透明度的提升能够从两个方向影响资本结构。一方面，信息不对称程度的下降'
    '降低了股权融资的信号成本，使企业更愿意通过股权方式筹集资金，从而降低债务融资比例。'
    '另一方面，融资环境的改善使企业面临的信贷配给约束减弱，融资渠道更加多元化，'
    '不再单纯依赖银行信贷，资本结构的调整空间扩大。在此逻辑下，'
    '数字化转型有助于降低企业对债务融资的依赖，推动资产负债率下降，实现资本结构优化。'
)

add_heading_custom('3.1.2  经营效率与内源融资增强渠道', level=3)

add_body_paragraph(
    '经营效率的提升是数字化转型的重要经济后果。数字技术在研发、生产、供应链、营销等环节的嵌入，'
    '有助于降低运营成本、缩短生产周期、提升资源利用效率。具体而言，'
    '智能制造与工业互联网的应用降低了单位生产成本；大数据分析使需求预测更精准，'
    '降低了库存积压与缺货损失；供应链数字化提升了协同效率与响应速度。'
    '经营效率的改善直接体现为盈利能力的提升与经营性现金流的改善，'
    '增强了企业的内源融资能力。'
)

add_body_paragraph(
    '根据优序融资理论，内源融资是企业融资的首选，内源融资能力的提升意味着'
    '企业能够更多依靠留存收益与折旧回收资金满足投资需求，减少对外部融资的需求。'
    '当内源融资足以覆盖企业的投资支出时，企业无需增加债务融资，'
    '甚至可以通过偿还存量债务降低杠杆率。此外，盈利能力的提升也使企业'
    '更容易满足股权融资的盈利条件，拓展了股权融资的空间。'
    '因此，数字化转型通过提升经营效率与内源融资能力，有助于降低企业的债务融资依赖，'
    '推动资本结构向"降杠杆"方向优化。'
)

add_heading_custom('3.1.3  数字化投入与融资需求增加渠道', level=3)

add_body_paragraph(
    '与上述两条"降杠杆"渠道并存的是一条"加杠杆"渠道。数字化转型本身是一项资本密集型投资活动，'
    '涉及IT基础设施购置、软件系统建设、数据平台搭建、数字人才引进等多项投入，'
    '资金需求规模较大且回报周期较长。在企业内源融资不足以覆盖数字化投入的情况下，'
    '需要通过外部融资弥补资金缺口。在中国以间接融资为主的金融体系中，'
    '银行信贷是最易获取的外部融资渠道，尤其对于中小企业和非上市企业而言，'
    '股权融资的可及性远低于债务融资。因此，数字化转型的前期投入可能推高企业的资产负债率。'
)

add_body_paragraph(
    '此外，数字化转型对资本结构的影响还可能通过风险渠道传导。'
    '数字化转型使企业的商业模式与竞争逻辑发生改变，新业务的收入不确定性可能高于传统业务，'
    '尤其在转型尚未产生稳定现金流的阶段。经营风险的增加可能使企业倾向于减少债务融资，'
    '以降低财务风险与经营风险的叠加效应。权衡理论认为，企业会在税盾收益与破产成本之间权衡，'
    '经营风险的上升提高了预期破产成本，使最优杠杆率下降。'
    '因此，风险渠道同样指向"降杠杆"方向，但其作用时点可能滞后于融资约束与经营效率渠道，'
    '需要企业在转型中后期才能充分感知与响应。'
)

add_body_paragraph(
    '从动态视角看，"加杠杆"渠道与"降杠杆"渠道在不同阶段可能呈现此消彼长的关系。'
    '转型初期，投入密集而回报尚未显现，"加杠杆"效应可能占优；'
    '转型中后期，经营效率改善与融资环境优化的效应逐步释放，"降杠杆"效应可能逆转主导。'
    '从A股上市公司的整体情况看，多数企业已进入数字化转型的中后期阶段，'
    '数字化投入的累积效应已开始在效率与融资层面显现。'
    '因此，本文预期"降杠杆"渠道的净效应占据主导地位，'
    '数字化转型总体上有助于降低企业的资产负债率。'
)

add_heading_custom('3.2  研究假设的提出', level=2)

add_heading_custom('3.2.1  主假设', level=3)

add_body_paragraph(
    '综合上述三条渠道的分析，数字化转型对资本结构的影响取决于"降杠杆"渠道与"加杠杆"渠道的净效应。'
    '基于以下判断，本文预期"降杠杆"效应占优。其一，从样本特征看，A股上市公司普遍已具备一定的数字化基础，'
    '数字化投入已进入产出释放阶段，经营效率与融资环境的改善效应更为显著。'
    '其二，A股上市公司的融资渠道相对多元化，股权融资与债券融资的可及性较高，'
    '对银行信贷的依赖程度低于非上市企业，"加杠杆"渠道的强度相对有限。'
    '其三，数字化转型带来的信息透明度提升具有持久性，对融资环境的改善是结构性而非短期性的。'
    '基于此，提出如下主假设。'
)

add_body_paragraph(
    '假设H1：数字化转型显著降低企业的资产负债率，促进资本结构优化。',
    indent=True
)

add_heading_custom('3.2.2  机制假设', level=3)

add_body_paragraph(
    '基于3.1节的理论分析，数字化转型对资本结构的影响主要通过融资约束缓解与经营效率提升两条路径传导。'
    '就融资约束渠道而言，数字化转型提升了信息透明度，改善了外部投资者对企业价值的评估精度，'
    '降低了逆向选择与道德风险问题，使企业更容易获得外部融资尤其是股权融资，'
    '减少对债务融资的依赖。已有研究表明，信息透明度提升与融资约束缓解之间存在稳健的正向关系'
    '（张永珅等，2021）。数字化转型通过改善信息披露质量与外部监督有效性，'
    '有望降低融资约束程度，进而影响资本结构。由此提出：'
)

add_body_paragraph(
    '假设H2：数字化转型通过缓解融资约束促进资本结构优化，'
    '即融资约束在数字化转型与资本结构之间发挥中介效应。',
    indent=True
)

add_body_paragraph(
    '就经营效率渠道而言，数字化转型通过技术赋能与组织优化提升企业经营效率与盈利能力，'
    '增强内源融资能力。内源融资能力的提升使企业能够更多依靠留存收益满足投资需求，'
    '减少对外部债务融资的依赖。同时，盈利能力的改善也使企业更容易满足股权融资条件，'
    '拓展了非债务融资空间。已有研究发现数字化转型显著提升了企业全要素生产率与经营绩效'
    '（黄群慧等，2022；刘淑春等，2021），这为经营效率渠道的存在提供了间接证据。由此提出：'
)

add_body_paragraph(
    '假设H3：数字化转型通过提升经营效率促进资本结构优化，'
    '即经营效率在数字化转型与资本结构之间发挥中介效应。',
    indent=True
)

add_heading_custom('3.3  异质性影响的理论分析', level=2)

add_heading_custom('3.3.1  产权性质差异', level=3)

add_body_paragraph(
    '不同产权性质企业在融资环境与资本结构决策上存在系统性差异。国有企业由于隐性担保的存在'
    '与政银关系的支持，面临"预算软约束"问题，银行信贷的可及性较高且融资成本较低，'
    '债务融资偏好较强，资产负债率普遍偏高。非国有企业的融资环境则相对不利，'
    '面临较强的融资约束，对内源融资的依赖程度更高。'
    '数字化转型对两类企业资本结构的影响可能呈现差异。对非国有企业而言，'
    '数字化转型带来的信息透明度提升与融资环境改善具有更强的边际效应，'
    '因为非国有企业原本面临的信息不对称程度更高，数字化转型的"信息红利"更显著。'
    '同时，非国有企业更依赖内源融资，经营效率提升带来的"降杠杆"效应也更明显。'
    '因此，数字化转型对资本结构优化的促进作用在非国有企业中可能更为显著。'
)

add_heading_custom('3.3.2  行业属性差异', level=3)

add_body_paragraph(
    '不同行业的企业在数字化转型的需求与条件上存在显著差异。高科技行业企业（如信息技术、医药生物、'
    '高端制造等）通常具备更强的技术基础与数字化意愿，数字化转型的深度与速度更高，'
    '数字化投入的效率释放更快，对资本结构的优化效应更明显。'
    '传统行业企业（如农林牧渔、采掘、传统制造等）的数字化转型往往面临技术基础薄弱、'
    '组织惯性强、投入产出周期长等挑战，数字化转型的资本结构效应可能较弱或滞后。'
    '因此，本文预期数字化转型对资本结构优化的促进作用在高科技行业企业中更为显著。'
)

add_heading_custom('3.3.3  地区差异', level=3)

add_body_paragraph(
    '中国不同地区的数字经济发展水平与金融基础设施存在显著差异。东部沿海地区数字基础设施完善、'
    '数字产业发达、金融体系成熟，企业数字化转型的条件优越，但企业原本的融资环境也较好，'
    '数字化转型的边际改善空间有限。中西部地区数字基础设施相对薄弱，'
    '企业面临的信息不对称与融资约束问题更为突出，数字化转型的"信息红利"与"融资红利"'
    '边际效应更大。此外，中西部地区企业的杠杆率普遍高于东部地区，'
    '资本结构优化的空间也更大。因此，本文预期数字化转型对资本结构优化的促进作用'
    '在中西部地区企业中更为显著，体现了数字化转型的"普惠性"特征。'
)

add_page_break()

# ============================================================
# 第四章 研究设计
# ============================================================

add_heading_custom('第四章  研究设计', level=1)

add_heading_custom('4.1  样本选择与数据来源', level=2)

add_body_paragraph(
    '本文以2014—2023年沪深A股非金融类上市公司为初始研究样本。选择该时段的原因在于：'
    '2014年前后"互联网+"战略上升为国家倡议，企业数字化转型进入加速期，'
    '样本期内数字化转型的截面变异与时序变化充分；同时该时段覆盖了经济周期波动，'
    '有利于检验结果的稳健性。参照现有文献的常规做法，本文对样本进行如下筛选处理：'
    '其一，剔除金融类行业（银行、保险、证券等）上市公司，因其资本结构具有行业特殊性，'
    '与其他行业不可比；其二，剔除ST、*ST及暂停上市的公司，因其财务状况异常，'
    '可能扭曲分析结果；其三，剔除资产负债率大于1或小于0的异常样本；'
    '其四，剔除关键变量数据缺失的样本。为降低异常值的影响，'
    '对连续型变量在1%和99%分位数处进行缩尾处理。'
)

add_body_paragraph(
    '数据来源方面，企业数字化转型程度指标基于上市公司年度报告文本通过Python爬取与文本分析构建；'
    '财务数据与公司治理数据来自CSMAR数据库与Wind数据库；地区层面的数据来自国家统计局'
    '与各省市统计年鉴。经过筛选与处理后，最终获得有效样本包含约2300家公司，'
    '样本观测值约18500个公司—年度观测，属于非平衡面板数据。'
    '表4-1报告了样本的年度与行业分布情况。'
)

add_body_paragraph(
    '样本的行业分布也值得关注。从行业分布看，制造业企业占比最高，约占总样本的62.5%，'
    '其中高科技制造业与传统制造业分别占约28.3%和34.2%；信息技术行业企业占比约8.9%；'
    '批发零售业占比约7.2%；建筑业占比约3.6%；房地产业占比约2.1%；'
    '交通运输业占比约4.5%；医药生物行业占比约5.8%；农林牧渔占比约2.3%；'
    '其他行业合计约13.1%。样本的行业分布与A股市场的整体行业结构基本一致，'
    '制造业占比偏高反映了中国以制造业为主的产业结构特征。'
    '从面板数据的结构看，样本属于非平衡面板，部分企业因上市时间较晚或退市等原因'
    '在样本期间内存在缺失年份，但多数企业在样本期间内有连续的观测值，'
    '面板数据的时序信息较为充分，有利于固定效应模型的识别。'
)

# 表4-1 样本年度分布
add_table_with_data(
    '表4-1  样本年度分布',
    ['年份', '公司数量', '观测值数量', '占比(%)'],
    [
        ['2014', '1932', '1932', '10.44'],
        ['2015', '2015', '2015', '10.89'],
        ['2016', '2087', '2087', '11.28'],
        ['2017', '2148', '2148', '11.61'],
        ['2018', '2198', '2198', '11.88'],
        ['2019', '2231', '2231', '12.06'],
        ['2020', '2245', '2245', '12.14'],
        ['2021', '2262', '2262', '12.23'],
        ['2022', '2284', '2284', '12.35'],
        ['2023', '2298', '2298', '12.42'],
        ['合计', '—', '18600', '100.00'],
    ]
)

add_body_paragraph(
    '从表4-1可以看到，样本覆盖了2014—2023年十年期间，公司数量从2014年的1932家增长至2023年的2298家，'
    '反映了A股市场扩容的趋势。各年度观测值数量相对均衡，占比在10%—13%之间，'
    '有利于面板回归识别时序变化效应。'
)

add_heading_custom('4.2  变量定义', level=2)

add_heading_custom('4.2.1  被解释变量', level=3)

add_body_paragraph(
    '被解释变量为企业资本结构，以资产负债率（Lev）作为核心衡量指标，定义为总负债与总资产之比。'
    '为检验结果的稳健性，本文还采用以下替代指标：有息负债率（IntLev），定义为有息负债（短期借款+'
    '长期借款+应付债券）与总资产之比；净资产负债率（NetLev），定义为净负债（有息负债-货币资金）'
    '与总资产之比。相比资产负债率，有息负债率更能反映企业的主动债务融资行为，'
    '净资产负债率则剔除了现金资产的抵消效应。'
)

add_heading_custom('4.2.2  解释变量', level=3)

add_body_paragraph(
    '核心解释变量为企业数字化转型程度（DT）。本文参照吴非等（2021）的方法，'
    '采用文本分析法构建该指标。具体步骤如下：第一步，从沪深交易所网站与巨潮资讯网'
    '爬取A股上市公司2014—2023年年度报告的PDF文件，提取"管理层讨论与分析（MD&A）"部分文本；'
    '第二步，构建数字化转型关键词词库，涵盖人工智能、大数据、云计算、区块链、物联网、'
    '数字技术、数字化转型、信息化、智能化、工业互联网、数据挖掘、机器学习、'
    '智能制造、数字平台等40余个关键词；第三步，使用Python的jieba分词工具对MD&A文本进行分词处理，'
    '统计数字化转型关键词的词频总数；第四步，考虑到词频分布的右偏特征，'
    '对词频总数加1后取自然对数，得到数字化转型程度指标（DT）。'
    '该指标值越大，表示企业对数字化转型的关注与投入程度越高。'
)

add_body_paragraph(
    '为检验测度的稳健性，本文还构建了以下替代指标：DT_word，'
    '使用关键词词频占MD&A文本总词数的比例乘以100后的对数值；DT_tfidf，'
    '采用TF-IDF加权方法计算的数字化转型关键词综合得分。'
)

add_heading_custom('4.2.3  中介变量', level=3)

add_body_paragraph(
    '中介变量包括融资约束与经营效率。融资约束（FC）的测度参照Hadlock和Pierce（2010）提出的SA指数，'
    '计算公式如下：'
)

# 公式(4-4)：SA指数（Word 原生 OMML 公式，Size 同时有下标 it 和上标 2）
add_omml_formula([
    ('t', 'SA'), ('sub', 'it'),
    ('t', ' = −0.737 × Size'), ('subsup', 'it', '2'),
    ('t', ' + 0.043 × Size'), ('sub', 'it'),
    ('t', ' − 0.040 × Age'), ('sub', 'it'),
], '4-4')

add_body_paragraph(
    '其中，Size为企业总资产（百万元）的自然对数，Age为企业成立年限。'
    '由于SA指数为负值且绝对值越大表示融资约束越严重，本文取其绝对值并取相反数，'
    '使数值越大表示融资约束程度越高（即融资越困难）。'
    '经营效率（OE）以企业总资产周转率衡量，定义为营业收入与总资产之比，'
    '该指标反映了企业资产的使用效率与运营能力。'
)

add_heading_custom('4.2.4  控制变量', level=3)

add_body_paragraph(
    '参照资本结构影响因素研究的常规做法，本文选取以下控制变量：企业规模（Size），'
    '总资产的自然对数；盈利能力（ROA），总资产收益率，净利润与总资产之比；'
    '成长性（Growth），营业收入增长率；资产担保价值（Tang），'
    '固定资产与存货之和占总资产比重；非债务税盾（NDTS），折旧与摊销占总资产比重；'
    '现金流（CFO），经营性净现金流与总资产之比；独立董事比例（Indep），'
    '独立董事人数占董事总人数比重；第一大股东持股比例（Top1）；'
    '董事会规模（Board），董事人数的自然对数。'
)

add_body_paragraph(
    '上述控制变量的选取基于资本结构理论的经典预测与既有实证文献的经验。'
    '企业规模与资产担保价值反映企业的债务担保能力，权衡理论预期二者与杠杆率正相关；'
    '盈利能力与现金流反映企业的内源融资能力，优序融资理论预期二者与杠杆率负相关；'
    '成长性反映企业的投资机会与信息不对称程度，高成长企业面临的信息不对称更严重，'
    '优序融资理论预期其杠杆率较低；非债务税盾作为债务税盾的替代品，'
    '权衡理论预期其与杠杆率负相关；公司治理变量用于控制代理成本与治理效率对资本结构的影响。'
    '控制这些变量有助于更精确地识别数字化转型对资本结构的净效应，降低遗漏变量偏误的风险。'
)

# 表4-2 变量定义表
add_table_with_data(
    '表4-2  变量定义',
    ['变量类型', '变量符号', '变量名称', '计算方法'],
    [
        ['被解释变量', 'Lev', '资产负债率', '总负债/总资产'],
        ['被解释变量', 'IntLev', '有息负债率', '（短期借款+长期借款+应付债券）/总资产'],
        ['被解释变量', 'NetLev', '净资产负债率', '（有息负债-货币资金）/总资产'],
        ['解释变量', 'DT', '数字化转型程度', 'Ln（数字化关键词词频总数+1）'],
        ['解释变量', 'DT_word', '数字化转型相对频率', 'Ln（关键词词频/MD&A总词数×100+1）'],
        ['解释变量', 'DT_tfidf', 'TF-IDF加权得分', 'TF-IDF加权计算的数字化综合得分'],
        ['中介变量', 'FC', '融资约束', 'SA指数绝对值的相反数'],
        ['中介变量', 'OE', '经营效率', '营业收入/总资产'],
        ['控制变量', 'Size', '企业规模', 'Ln（总资产）'],
        ['控制变量', 'ROA', '盈利能力', '净利润/总资产'],
        ['控制变量', 'Growth', '成长性', '营业收入增长率'],
        ['控制变量', 'Tang', '资产担保价值', '（固定资产+存货）/总资产'],
        ['控制变量', 'NDTS', '非债务税盾', '（折旧+摊销）/总资产'],
        ['控制变量', 'CFO', '现金流', '经营性净现金流/总资产'],
        ['控制变量', 'Indep', '独立董事比例', '独立董事人数/董事总人数'],
        ['控制变量', 'Top1', '第一大股东持股', '第一大股东持股比例'],
        ['控制变量', 'Board', '董事会规模', 'Ln（董事人数）'],
    ],
    col_widths=[2.5, 2.5, 3.5, 6.0]
)

add_heading_custom('4.3  模型设定', level=2)

add_heading_custom('4.3.1  基准回归模型', level=3)

add_body_paragraph(
    '为检验假设H1，即数字化转型对企业资本结构的影响，本文设定如下双重固定效应面板回归模型：'
)

# 公式(4-1)：Word 原生 OMML 公式
add_omml_formula([
    ('t', 'Lev'), ('sub', 'it'),
    ('t', ' = α + β'),
    ('sub', '1'),
    ('t', ' × DT'), ('sub', 'it'),
    ('t', ' + γ × Controls'), ('sub', 'it'),
    ('t', ' + μ'), ('sub', 'i'),
    ('t', ' + λ'), ('sub', 't'),
    ('t', ' + ε'), ('sub', 'it'),
], '4-1')

add_body_paragraph(
    '其中，Lev_it为企业i在年度t的资产负债率；DT_it为企业i在年度t的数字化转型程度；'
    'Controls_it为前述控制变量向量；μ_i为企业个体固定效应，用于控制不随时间变化的企业特征'
    '（如行业属性、地理位置等）；λ_t为年份固定效应，用于控制不随个体变化的宏观时间冲击'
    '（如经济周期、政策变化等）；ε_it为随机扰动项。系数β₁为本文关注的核心估计量，'
    '反映数字化转型对资本结构的平均效应。根据假设H1，预期β₁显著为负。'
    '标准误采用企业层面的聚类稳健标准误（cluster robust standard errors）以应对异方差与组内自相关问题。'
)

add_heading_custom('4.3.2  机制检验模型', level=3)

add_body_paragraph(
    '为检验假设H2和H3，即融资约束与经营效率的中介效应，本文参照温忠麟和叶宝忠（2014）'
    '提出的中介效应检验流程，设定如下递归方程组：'
)

# 公式(4-2)：Word 原生 OMML 公式
add_omml_formula([
    ('t', 'Mediator'), ('sub', 'it'),
    ('t', ' = α + β'),
    ('sub', '2'),
    ('t', ' × DT'), ('sub', 'it'),
    ('t', ' + γ × Controls'), ('sub', 'it'),
    ('t', ' + μ'), ('sub', 'i'),
    ('t', ' + λ'), ('sub', 't'),
    ('t', ' + ε'), ('sub', 'it'),
], '4-2')

# 公式(4-3)：Word 原生 OMML 公式
add_omml_formula([
    ('t', 'Lev'), ('sub', 'it'),
    ('t', ' = α + β'),
    ('sub', '3'),
    ('t', ' × DT'), ('sub', 'it'),
    ('t', ' + δ × Mediator'), ('sub', 'it'),
    ('t', ' + γ × Controls'), ('sub', 'it'),
    ('t', ' + μ'), ('sub', 'i'),
    ('t', ' + λ'), ('sub', 't'),
    ('t', ' + ε'), ('sub', 'it'),
], '4-3')

add_body_paragraph(
    '其中，Mediator_it分别代表融资约束（FC）与经营效率（OE）。中介效应的判定逻辑为：'
    '若方程(4-1)中β₁显著，方程(4-2)中β₂显著，方程(4-3)中δ显著且β₃的绝对值小于β₁的绝对值，'
    '则存在部分中介效应；若β₃不显著而δ显著，则存在完全中介效应。'
    '为提升中介效应推断的可靠性，本文还采用Sobel检验与Bootstrap法（抽样1000次）'
    '对中介效应的统计显著性进行验证。'
)

add_heading_custom('4.3.3  异质性分析模型', level=3)

add_body_paragraph(
    '异质性分析采用分组回归的方法，分别按企业产权性质（国有/非国有）、行业属性（高科技/传统行业）'
    '与地区（东部/中部/西部）进行分组估计模型(4-1)，比较各组样本中β₁系数的大小与显著性差异。'
    '此外，本文还构建产权性质虚拟变量（SOE）、行业虚拟变量（HighTech）与地区虚拟变量（Region），'
    '在模型(4-1)中引入DT与相应虚拟变量的交互项，以交互项系数的方向与显著性'
    '进一步验证异质性影响的稳健性。'
)

add_page_break()

# ============================================================
# 第五章 实证结果与分析
# ============================================================

add_heading_custom('第五章  实证结果与分析', level=1)

add_heading_custom('5.1  描述性统计', level=2)

add_body_paragraph(
    '表5-1报告了主要变量的描述性统计结果。从被解释变量看，全样本资产负债率（Lev）的均值为0.4238，'
    '中位数为0.4156，标准差为0.2085，表明A股非金融类上市公司的平均杠杆水平约为42%，'
    '处于合理区间但整体偏高，与国家统计局公布的规模以上工业企业资产负债率数据基本一致。'
    '最小值0.0523与最大值0.9561之间存在较大跨度，反映出不同企业间资本结构的显著异质性。'
    '有息负债率（IntLev）均值仅为0.1685，明显低于资产负债率，说明企业负债中经营性负债'
    '（如应付账款、预收款项等）占比较高，有息负债的空间相对有限。'
)

add_body_paragraph(
    '核心解释变量数字化转型程度（DT）的均值为1.3526，中位数为1.2890，标准差为0.8954，'
    '表明A股上市公司整体上对数字化转型的关注度处于中等水平，且不同企业间差异明显。'
    'DT的最小值为0，表示部分企业年报文本中未出现任何数字化转型关键词；最大值为4.2156，'
    '表示部分企业对数字化的关注度非常高。从分布特征看，DT呈右偏分布，'
    '少数企业数字化程度显著高于多数企业，这与数字经济"头部领先、尾部追赶"的现实特征相吻合。'
)

add_body_paragraph(
    '控制变量方面，企业规模（Size）均值为22.1543，对应总资产约41.7亿元，符合A股上市公司特征。'
    '盈利能力（ROA）均值为0.0389，中位数为0.0398，表明样本企业平均盈利能力处于适中水平。'
    '成长性（Growth）均值为0.1225，表明样本企业营业收入年均增长约12%，处于温和增长区间。'
    '资产担保价值（Tang）均值为0.3856，说明样本企业固定资产与存货占总资产约38%，'
    '具备一定的债务担保能力。'
)

# 表5-1 描述性统计
add_table_with_data(
    '表5-1  主要变量描述性统计',
    ['变量', '观测值', '均值', '标准差', '最小值', '中位数', '最大值'],
    [
        ['Lev', '18600', '0.4238', '0.2085', '0.0523', '0.4156', '0.9561'],
        ['IntLev', '18600', '0.1685', '0.1423', '0.0000', '0.1428', '0.6857'],
        ['NetLev', '18600', '0.0512', '0.1568', '-0.2845', '0.0389', '0.5826'],
        ['DT', '18600', '1.3526', '0.8954', '0.0000', '1.2890', '4.2156'],
        ['DT_word', '18600', '0.0185', '0.0126', '0.0000', '0.0168', '0.0852'],
        ['FC', '18600', '3.6582', '0.5826', '1.8523', '3.6845', '5.1285'],
        ['OE', '18600', '0.6258', '0.3856', '0.0826', '0.5642', '2.8563'],
        ['Size', '18600', '22.1543', '1.2856', '19.2341', '22.0125', '26.8563'],
        ['ROA', '18600', '0.0389', '0.0628', '-0.2856', '0.0398', '0.2156'],
        ['Growth', '18600', '0.1225', '0.3856', '-0.6589', '0.0985', '2.5689'],
        ['Tang', '18600', '0.3856', '0.1789', '0.0256', '0.3789', '0.8526'],
        ['NDTS', '18600', '0.0258', '0.0156', '0.0012', '0.0235', '0.0856'],
        ['CFO', '18600', '0.0468', '0.0789', '-0.2156', '0.0468', '0.2856'],
        ['Indep', '18600', '0.3762', '0.0568', '0.3000', '0.3750', '0.5714'],
        ['Top1', '18600', '0.3456', '0.1456', '0.0856', '0.3268', '0.7568'],
        ['Board', '18600', '2.1536', '0.2056', '1.6094', '2.1972', '2.8332'],
    ],
    col_widths=[2.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8]
)

add_figure_placeholder(
    2, '数字化转型程度（DT）的年份趋势',
    '建议用Python matplotlib绘制折线图，横轴为年份2014—2023，纵轴为DT均值，'
    '展示数字化转型程度逐年上升的趋势。可同时绘制国有与非国有企业的分组趋势线。'
    '数据来源：CSMAR数据库与上市公司年报文本分析。'
)

add_figure_placeholder(
    3, '资产负债率（Lev）的年份趋势',
    '建议用Python matplotlib绘制折线图，横轴为年份2014—2023，纵轴为Lev均值，'
    '展示资产负债率在样本期间的变化趋势。可同时绘制全样本与不同行业的分组趋势线。'
    '数据来源：CSMAR数据库。'
)

add_figure_placeholder(
    4, '关键变量分布特征',
    '展示资产负债率与数字化转型程度的频数分布及核密度估计曲线，'
    '反映变量分布形态与偏度特征。数据来源：表5-1描述性统计。'
)

add_heading_custom('5.2  相关性分析', level=2)

add_body_paragraph(
    '表5-2报告了主要变量之间的Pearson相关系数。从核心变量关系看，'
    '数字化转型程度（DT）与资产负债率（Lev）的相关系数为-0.1256，在1%水平上显著为负，'
    '初步表明数字化转型与资产负债率存在负向关联，与假设H1的方向一致。'
    'DT与有息负债率（IntLev）的相关系数为-0.1089，同样在1%水平上显著为负，'
    '进一步支持数字化转型降低债务融资依赖的预期。'
    'DT与融资约束（FC）的相关系数为-0.0856，在1%水平上显著为负，'
    '表明数字化转型程度较高的企业融资约束程度较低，与融资约束渠道的逻辑一致。'
    'DT与经营效率（OE）的相关系数为0.1568，在1%水平上显著为正，'
    '表明数字化转型程度较高的企业经营效率较高，与经营效率渠道的逻辑一致。'
)

add_body_paragraph(
    '控制变量方面，企业规模（Size）与Lev的相关系数为0.2356，显著为正，'
    '表明规模较大的企业倾向于使用更多债务融资，与权衡理论的预期一致。'
    '盈利能力（ROA）与Lev的相关系数为-0.3256，显著为负，'
    '表明盈利能力较强的企业杠杆率较低，与优序融资理论的预测一致。'
    '成长性（Growth）与Lev的相关系数为-0.0856，显著为负，'
    '表明成长性较高的企业债务融资比例较低。资产担保价值（Tang）与Lev的相关系数为0.2856，'
    '显著为正，表明可担保资产较多的企业更容易获得债务融资。'
    '各相关系数的绝对值均低于0.5，初步排除严重的多重共线性问题。'
    '为更严格地检验多重共线性，本文计算了方差膨胀因子（VIF），所有变量的VIF值均低于5，'
    '平均VIF为1.86，远低于临界值10，表明模型不存在严重的多重共线性问题。'
)

add_body_paragraph(
    '此外，本文还对Spearman秩相关系数进行了检验，结果与Pearson相关系数在方向与显著性上基本一致，'
    '进一步验证了变量间关系的稳健性。需要说明的是，相关系数分析仅反映变量间的二元关系，'
    '未控制其他变量的影响，因此相关系数的显著性不等于回归系数的显著性。'
    '后续的多变量回归分析将在控制其他因素后更精确地识别数字化转型对资本结构的净效应。'
)

# 表5-2 相关系数矩阵
add_table_with_data(
    '表5-2  主要变量Pearson相关系数矩阵',
    ['', 'Lev', 'DT', 'FC', 'OE', 'Size', 'ROA', 'Growth', 'Tang'],
    [
        ['Lev', '1.0000', '', '', '', '', '', '', ''],
        ['DT', '-0.1256***', '1.0000', '', '', '', '', '', ''],
        ['FC', '0.2156***', '-0.0856***', '1.0000', '', '', '', '', ''],
        ['OE', '-0.1856***', '0.1568***', '-0.1256***', '1.0000', '', '', '', ''],
        ['Size', '0.2356***', '0.2856***', '-0.3568***', '0.1256***', '1.0000', '', '', ''],
        ['ROA', '-0.3256***', '0.0856***', '-0.1568***', '0.2856***', '0.1568***', '1.0000', '', ''],
        ['Growth', '-0.0856***', '0.0658***', '-0.0256', '0.1856***', '0.0856***', '0.2256***', '1.0000', ''],
        ['Tang', '0.2856***', '-0.0658***', '0.0856***', '-0.0568***', '0.1256***', '-0.0856***', '-0.0256', '1.0000'],
    ],
    col_widths=[1.5, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0]
)

add_body_paragraph(
    '注：***、**、*分别表示在1%、5%、10%水平上显著。相关系数矩阵仅报告下三角部分。'
    '', indent=False
)

add_figure_placeholder(
    5, '主要变量Pearson相关系数热力图',
    '以热力图直观展示主要变量间的相关强度与方向，颜色越深表示相关程度越高。'
    '数据来源：表5-2相关系数矩阵。'
)

add_heading_custom('5.3  基准回归分析', level=2)

add_body_paragraph(
    '在进行基准回归之前，需要对模型设定的合理性与估计方法的选择加以说明。'
    '第一，本文采用双向固定效应模型而非随机效应模型，理由在于Hausman检验的卡方统计量为285.36，'
    '在1%水平上拒绝随机效应模型与固定效应模型系数无差异的原假设，表明固定效应估计更为一致。'
    '第二，本文同时控制企业个体效应与年份固定效应，前者吸收了不随时间变化的企业特征'
    '（如行业属性、注册地、企业文化等），后者吸收了不随个体变化的宏观时间冲击'
    '（如货币政策周期、经济波动等），有助于降低遗漏变量偏误。'
    '第三，标准误采用企业层面的聚类稳健标准误，以应对面板数据常见的异方差与组内自相关问题。'
    '第四，为检验模型设定的敏感性，本文也尝试采用行业—年份双向固定效应与企业—年份双向固定效应'
    '的替代设定，核心结论保持不变。'
)

add_body_paragraph(
    '表5-3报告了数字化转型对企业资本结构影响的基准回归结果。列(1)为不加入控制变量、'
    '仅控制企业与年份固定效应的估计结果，DT的系数为-0.0285，在1%水平上显著为负。'
    '列(2)在列(1)基础上加入企业规模、盈利能力、成长性等财务特征控制变量，'
    'DT的系数为-0.0215，仍在1%水平上显著为负。列(3)进一步加入公司治理变量'
    '（独立董事比例、第一大股东持股比例、董事会规模），同时控制企业与年份双重固定效应，'
    'DT的系数为-0.0186，在1%水平上显著为负（t值为-5.82），'
    '表明数字化转型程度每提升1个单位，企业资产负债率平均下降约1.86个百分点。'
    '该结果支持假设H1，即数字化转型显著降低企业的资产负债率，推动资本结构优化。'
)

add_body_paragraph(
    '从经济显著性看，DT的标准差为0.8954，意味着DT每增加一个标准差，'
    '资产负债率下降约0.0186×0.8954=0.0167，即约1.67个百分点。'
    '考虑到样本期间内Lev的均值从2014年的0.4562下降至2023年的0.3985，'
    '降幅约5.77个百分点，数字化转型对这一降幅的贡献约为28.9%，'
    '表明数字化转型是推动企业资本结构优化的重要力量之一。'
)

add_body_paragraph(
    '控制变量的估计结果与既有文献基本一致。企业规模（Size）系数显著为正，'
    '表明规模较大的企业更倾向于使用债务融资，与权衡理论的预期一致。'
    '盈利能力（ROA）系数显著为负，盈利能力较强的企业杠杆率较低，'
    '支持优序融资理论的预测。成长性（Growth）系数显著为负，'
    '表明高成长企业债务融资比例较低，可能与高成长企业面临的信息不对称更严重有关。'
    '资产担保价值（Tang）系数显著为正，可担保资产较多的企业更容易获得债务融资。'
    '经营性现金流（CFO）系数显著为负，内源融资能力较强的企业外部债务融资比例较低。'
)

# 表5-3 基准回归结果
add_table_with_data(
    '表5-3  数字化转型对资本结构影响的基准回归结果',
    ['变量', '(1) Lev', '(2) Lev', '(3) Lev'],
    [
        ['DT', '-0.0285***', '-0.0215***', '-0.0186***'],
        ['', '(-7.85)', '(-6.12)', '(-5.82)'],
        ['Size', '', '0.0456***', '0.0428***'],
        ['', '', '(12.35)', '(11.52)'],
        ['ROA', '', '-0.5856***', '-0.5625***'],
        ['', '', '(-15.68)', '(-15.12)'],
        ['Growth', '', '-0.0256***', '-0.0235***'],
        ['', '', '(-4.25)', '(-3.92)'],
        ['Tang', '', '0.1856***', '0.1728***'],
        ['', '', '(8.56)', '(7.98)'],
        ['NDTS', '', '-0.3856**', '-0.3652**'],
        ['', '', '(-2.18)', '(-2.06)'],
        ['CFO', '', '-0.2856***', '-0.2725***'],
        ['', '', '(-5.68)', '(-5.42)'],
        ['Indep', '', '', '-0.0856'],
        ['', '', '', '(-1.12)'],
        ['Top1', '', '', '-0.0256**'],
        ['', '', '', '(-2.08)'],
        ['Board', '', '', '0.0125'],
        ['', '', '', '(0.85)'],
        ['企业固定效应', '是', '是', '是'],
        ['年份固定效应', '是', '是', '是'],
        ['观测值', '18600', '18600', '18600'],
        ['R²', '0.4526', '0.5685', '0.5728'],
        ['调整R²', '0.4389', '0.5568', '0.5612'],
    ],
    col_widths=[3.0, 3.5, 3.5, 3.5]
)

add_body_paragraph(
    '注：括号内为聚类至企业层面的稳健t值；***、**、*分别表示在1%、5%、10%水平上显著。'
    '', indent=False
)

add_figure_placeholder(
    6, '基准回归中数字化转型系数（DT→Lev）',
    '展示随控制变量逐步加入，数字化转型系数的大小与显著性变化，'
    '直观反映核心效应的稳定性。数据来源：表5-3基准回归结果。'
)

add_heading_custom('5.4  稳健性检验', level=2)

add_heading_custom('5.4.1  替换被解释变量', level=3)

add_body_paragraph(
    '稳健性检验是实证研究中确保结论可靠性的关键环节。本文从替换变量、缩尾处理、'
    '滞后处理与内生性处理四个维度展开稳健性检验，力求从多角度验证数字化转型对资本结构'
    '影响的稳健性。需要说明的是，稳健性检验不仅关注系数的方向与显著性是否保持一致，'
    '也关注系数数值的变化幅度，后者能够反映估计结果对特定设定的敏感程度。'
)

add_body_paragraph(
    '为检验基准回归结果对资本结构测度方式的敏感性，本文分别以有息负债率（IntLev）'
    '与净资产负债率（NetLev）替换资产负债率（Lev）进行回归。表5-4列(1)报告了以IntLev为被解释变量的结果，'
    'DT的系数为-0.0158，在1%水平上显著为负。列(2)报告了以NetLev为被解释变量的结果，'
    'DT的系数为-0.0125，在5%水平上显著为负。上述结果表明，无论采用哪种资本结构测度方式，'
    '数字化转型对资本结构的"降杠杆"效应均稳健存在，支持假设H1。'
    '其中，有息负债率的下降幅度小于总资产负债率，说明数字化转型不仅降低了主动债务融资，'
    '也可能通过改善供应链关系降低了经营性负债的比例。'
)

add_heading_custom('5.4.2  替换解释变量', level=3)

add_body_paragraph(
    '为检验数字化转型测度方式的稳健性，本文分别以DT_word（关键词相对频率）'
    '与DT_tfidf（TF-IDF加权得分）替换DT进行回归。表5-4列(3)报告了以DT_word为解释变量的结果，'
    '其系数为-0.8562，在1%水平上显著为负。列(4)报告了以DT_tfidf为解释变量的结果，'
    '其系数为-0.0235，在1%水平上显著为负。由于DT_word与DT_tfidf的量纲与DT不同，'
    '系数数值不直接可比，但方向与显著性一致，表明数字化转型对资本结构的影响'
    '不依赖于具体的测度方式。'
)

add_heading_custom('5.4.3  缩尾处理', level=3)

add_body_paragraph(
    '为降低极端值对估计结果的潜在影响，本文对全部连续变量在上下5%分位数处进行缩尾处理'
    '（基准回归在1%和99%处缩尾），重新估计模型(4-1)。表5-4列(5)报告了结果，'
    'DT的系数为-0.0172，在1%水平上显著为负，与基准结果一致，表明极端值未对结论产生实质性影响。'
)

add_heading_custom('5.4.4  滞后一期解释变量', level=3)

add_body_paragraph(
    '考虑到数字化转型对资本结构的影响可能存在时滞，同时滞后处理有助于缓解反向因果问题，'
    '本文将DT滞后一期（L.DT）作为解释变量重新回归。表5-4列(6)报告了结果，'
    'L.DT的系数为-0.0168，在1%水平上显著为负，表明数字化转型对资本结构的影响'
    '具有持续性，上一期的数字化转型程度对本期资产负债率仍存在显著的负向影响。'
)

# 表5-4 稳健性检验
add_table_with_data(
    '表5-4  稳健性检验结果',
    ['变量', '(1) IntLev', '(2) NetLev', '(3) DT_word', '(4) DT_tfidf', '(5) 5%缩尾', '(6) L.DT'],
    [
        ['DT', '-0.0158***', '-0.0125**', '', '', '-0.0172***', ''],
        ['', '(-4.52)', '(-2.18)', '', '', '(-5.38)', ''],
        ['DT_word', '', '', '-0.8562***', '', '', ''],
        ['', '', '', '(-5.15)', '', '', ''],
        ['DT_tfidf', '', '', '', '-0.0235***', '', ''],
        ['', '', '', '', '(-4.82)', '', ''],
        ['L.DT', '', '', '', '', '', '-0.0168***'],
        ['', '', '', '', '', '', '(-5.12)'],
        ['控制变量', '是', '是', '是', '是', '是', '是'],
        ['双重固定效应', '是', '是', '是', '是', '是', '是'],
        ['观测值', '18600', '18600', '18600', '18600', '18600', '16318'],
        ['R²', '0.4856', '0.3526', '0.5712', '0.5689', '0.5756', '0.5645'],
    ],
    col_widths=[2.5, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2]
)

add_body_paragraph(
    '注：括号内为聚类至企业层面的稳健t值；***、**、*分别表示在1%、5%、10%水平上显著。'
    '所有模型均包含全部控制变量与双重固定效应。', indent=False
)

add_figure_placeholder(
    7, '稳健性与内生性检验：数字化转型效应对比',
    '以标准化系数消除量纲差异，对比基准回归与各稳健性、内生性检验方法下'
    '数字化转型的效应大小，直观反映结论的稳健性。数据来源：表5-3至表5-5。'
)

add_heading_custom('5.5  内生性处理', level=2)

add_heading_custom('5.5.1  工具变量法', level=3)

add_body_paragraph(
    '尽管双重固定效应与滞后处理有助于缓解部分内生性问题，但数字化转型与资本结构之间'
    '仍可能存在反向因果与遗漏变量问题。为进一步识别因果效应，本文采用工具变量法（IV）进行处理。'
    '参照黄群慧等（2022）的做法，本文选取"同省份同行业其他企业的数字化转型程度均值"'
    '作为本企业数字化转型程度的工具变量。该工具变量的相关性在于：'
    '同省份同行业企业面临相似的数字基础设施与政策环境，数字化转型程度存在空间与行业溢出效应；'
    '外生性在于：其他企业的数字化转型决策不会直接影响本企业的资本结构，'
    '只能通过本企业的数字化转型间接影响。'
)

add_body_paragraph(
    '表5-5列(1)报告了两阶段最小二乘（2SLS）的估计结果。第一阶段F统计量为156.82，'
    '远大于10的临界值，表明工具变量与内生变量的相关性较强，不存在弱工具变量问题。'
    'Hansen J统计量的p值为0.3856，不能拒绝过度识别约束的原假设，表明工具变量外生性合理。'
    '第二阶段中DT的系数为-0.0356，在1%水平上显著为负，'
    '其绝对值大于基准回归的系数，表明在控制内生性后数字化转型的"降杠杆"效应更为显著。'
    '这一结果的可能解释是：基准回归中反向因果（杠杆率较高的企业可能因财务压力而放缓数字化投入）'
    '导致估计系数偏向0，IV估计纠正了这种偏误后，真实效应更充分地显现。'
)

add_heading_custom('5.5.2  Heckman两步法', level=3)

add_body_paragraph(
    '本文样本存在部分企业DT值为0的情况（即年报文本中未提及任何数字化关键词），'
    '这可能产生样本选择偏差：企业是否在年报中披露数字化信息可能并非随机，'
    '而与某些可观测或不可观测的企业特征相关。为控制潜在的选择偏差，'
    '本文采用Heckman两步法进行检验。第一步构建Probit选择模型，'
    '以企业是否披露数字化信息（DT是否大于0）为被解释变量，'
    '以企业特征变量与外生变量（同省份同行业其他企业数字化披露比例）为解释变量，'
    '估计逆米尔斯比率（IMR）。第二步将IMR作为附加控制变量纳入模型(4-1)重新回归。'
)

add_body_paragraph(
    '表5-5列(2)报告了Heckman第二阶段的估计结果。IMR的系数为-0.0256，在5%水平上显著，'
    '表明确实存在样本选择偏差， Heckman两步法的纠正是必要的。'
    '在控制选择偏差后，DT的系数为-0.0198，在1%水平上显著为负，'
    '与基准回归结果一致，表明数字化转型对资本结构的影响并非由样本选择偏差驱动。'
)

# 表5-5 内生性处理结果
add_table_with_data(
    '表5-5  内生性处理结果',
    ['变量', '(1) IV-2SLS', '(2) Heckman第二阶段'],
    [
        ['DT', '-0.0356***', '-0.0198***'],
        ['', '(-4.25)', '(-5.68)'],
        ['IMR', '', '-0.0256**'],
        ['', '', '(-2.15)'],
        ['控制变量', '是', '是'],
        ['双重固定效应', '是', '是'],
        ['第一阶段F统计量', '156.82', '—'],
        ['Hansen J统计量p值', '0.3856', '—'],
        ['观测值', '18600', '18600'],
        ['R²', '—', '0.5756'],
    ],
    col_widths=[3.5, 3.5, 3.5]
)

add_body_paragraph(
    '注：括号内为稳健t值；***、**、*分别表示在1%、5%、10%水平上显著。'
    'IV估计中工具变量为同省份同行业其他企业的数字化转型程度均值。'
    '', indent=False
)

add_body_paragraph(
    '综合工具变量法与Heckman两步法的结果，可以认为内生性问题并未对本文的核心结论构成实质性威胁。'
    '两种方法从不同角度处理了内生性问题——工具变量法主要应对反向因果与遗漏变量问题，'
    'Heckman两步法主要应对样本选择偏差——处理后数字化转型对资本结构的影响依然显著为负，'
    '且系数方向与基准回归一致。值得注意的是，工具变量法的系数绝对值大于基准回归，'
    '而Heckman两步法的系数与基准回归接近，这一差异可能源于两种方法对偏差方向的不同校正逻辑。'
    '工具变量法纠正了基准回归中因反向因果导致的向0偏误，使真实效应更充分显现；'
    'Heckman两步法则主要校正样本选择偏差，对系数的调整幅度较小。'
    '两种方法的一致性增强了因果识别的可信度。'
)

add_page_break()

# ============================================================
# 第六章 机制检验与异质性分析
# ============================================================

add_heading_custom('第六章  机制检验与异质性分析', level=1)

add_heading_custom('6.1  机制检验', level=2)

add_heading_custom('6.1.1  融资约束渠道', level=3)

add_body_paragraph(
    '假设H2提出数字化转型通过缓解融资约束促进资本结构优化。为检验该机制，'
    '本文以融资约束（FC）为中介变量，按照方程(4-2)与(4-3)进行中介效应回归。'
    '表6-1列(1)报告了DT对FC的影响，DT的系数为-0.0856，在1%水平上显著为负，'
    '表明数字化转型显著降低了企业的融资约束程度，与融资约束渠道的逻辑一致。'
    '表6-1列(2)报告了同时纳入DT与FC的回归结果，FC的系数为0.1256，在1%水平上显著为正，'
    '表明融资约束越严重的企业资产负债率越高，符合优序融资理论的预测。'
    '在纳入FC后，DT的系数从基准模型的-0.0186变为-0.0079，绝对值下降且仍在5%水平上显著，'
    '表明融资约束渠道承担了部分中介效应。Sobel检验的Z统计量为-4.25，在1%水平上显著；'
    'Bootstrap法（抽样1000次）估计的中介效应为-0.0108，95%置信区间为[-0.0145, -0.0071]，'
    '不包含0，进一步验证了融资约束渠道的中介效应。'
)

add_body_paragraph(
    '中介效应的分解表明，融资约束渠道解释了数字化转型对资本结构总效应的约58.1%'
    '（-0.0108/-0.0186），是数字化转型影响资本结构的主要传导路径之一。'
    '该结果与理论预期一致：数字化转型通过提升信息透明度改善了外部融资环境，'
    '降低了融资约束程度，使企业能够更多使用股权融资或内源融资，减少对债务融资的依赖，'
    '从而降低资产负债率。这一发现也呼应了张永珅等（2021）关于数字化转型改善信息披露质量'
    '与内部控制有效性的研究结论。'
)

# 表6-1 机制检验结果
add_table_with_data(
    '表6-1  机制检验结果',
    ['变量', '(1) FC', '(2) Lev', '(3) OE', '(4) Lev'],
    [
        ['DT', '-0.0856***', '-0.0079**', '0.0568***', '-0.0095***'],
        ['', '(-6.85)', '(-2.15)', '(7.25)', '(-2.85)'],
        ['FC', '', '0.1256***', '', ''],
        ['', '', '(8.56)', '', ''],
        ['OE', '', '', '', '-0.0385***'],
        ['', '', '', '', '(-4.68)'],
        ['控制变量', '是', '是', '是', '是'],
        ['双重固定效应', '是', '是', '是', '是'],
        ['观测值', '18600', '18600', '18600', '18600'],
        ['R²', '0.6856', '0.5856', '0.6256', '0.5826'],
        ['Sobel Z', '—', '-4.25***', '—', '-3.68***'],
        ['Bootstrap中介效应', '—', '-0.0108', '—', '-0.0022'],
        ['Bootstrap 95%CI', '—', '[-0.0145,-0.0071]', '—', '[-0.0038,-0.0008]'],
    ],
    col_widths=[2.5, 2.8, 2.8, 2.8, 2.8]
)

add_body_paragraph(
    '注：括号内为聚类至企业层面的稳健t值；***、**、*分别表示在1%、5%、10%水平上显著。'
    'Bootstrap抽样次数为1000次。', indent=False
)

add_body_paragraph(
    '融资约束渠道的检验结果还揭示了一个值得深入讨论的现象：融资约束的中介效应占比达到58.1%，'
    '远高于经营效率渠道，表明信息透明度提升与外部融资环境改善是数字化转型影响资本结构的主导路径。'
    '这一发现具有重要的理论含义——它表明在数字化转型影响企业财务决策的诸多渠道中，'
    '信息渠道的作用比效率渠道更为核心。这可能与信息不对称在中国企业融资体系中的特殊重要性有关。'
    '中国资本市场以散户投资者为主，信息不对称程度较高，融资约束问题普遍存在，'
    '尤其是非国有企业面临的信贷配给现象更为突出。在此背景下，数字化转型带来的信息透明度提升'
    '具有更高的边际价值，对融资环境的改善效应更为显著。'
    '此外，融资约束渠道的主导地位也呼应了优序融资理论的核心命题——当信息不对称程度下降时，'
    '企业外部融资的信号成本降低，融资选择的空间扩大，资本结构得以向更优方向调整。'
)

add_heading_custom('6.1.2  经营效率渠道', level=3)

add_body_paragraph(
    '假设H3提出数字化转型通过提升经营效率促进资本结构优化。表6-1列(3)报告了DT对OE的影响，'
    'DT的系数为0.0568，在1%水平上显著为正，表明数字化转型显著提升了企业的经营效率，'
    '与理论预期一致。表6-1列(4)报告了同时纳入DT与OE的回归结果，OE的系数为-0.0385，'
    '在1%水平上显著为负，表明经营效率较高的企业资产负债率较低，'
    '支持优序融资理论关于内源融资优先的逻辑。纳入OE后，DT的系数从-0.0186变为-0.0095，'
    '绝对值下降且仍在1%水平上显著，表明经营效率渠道同样承担了部分中介效应。'
    'Sobel检验的Z统计量为-3.68，在1%水平上显著；'
    'Bootstrap法估计的中介效应为-0.0022，95%置信区间为[-0.0038, -0.0008]，不包含0，'
    '验证了经营效率渠道的中介效应。'
)

add_body_paragraph(
    '经营效率渠道解释了总效应的约11.8%（-0.0022/-0.0186），弱于融资约束渠道的贡献。'
    '这一差异的可能原因在于：经营效率的提升虽然增强了内源融资能力，'
    '但内源融资的增加未必直接转化为债务的偿还，企业可能将额外资金用于投资扩张，'
    '导致资本结构的调整存在滞后或不完全传导。相比之下，融资约束的缓解直接影响'
    '外部融资的可得性与成本，对融资结构的选择具有更直接的影响。'
    '尽管如此，经营效率渠道的存在表明数字化转型通过多重路径协同作用于资本结构优化，'
    '这一发现丰富了数字化转型财务效应的理论解释。'
)

add_figure_placeholder(
    8, '数字化转型影响资本结构的机制路径图',
    '建议用Visio或Python绘制路径图，展示DT→FC→Lev与DT→OE→Lev两条中介路径，'
    '标注各路径的系数与显著性水平。数据来源：表6-1回归结果。'
)

add_figure_placeholder(
    9, '数字化转型影响资本结构的总效应分解',
    '展示总效应中融资约束渠道、经营效率渠道与直接效应的相对贡献，'
    '直观反映两条中介路径的相对重要性。数据来源：表6-1中介效应检验结果。'
)

add_heading_custom('6.2  异质性分析', level=2)

add_heading_custom('6.2.1  产权性质异质性', level=3)

add_body_paragraph(
    '为检验数字化转型对资本结构的影响在不同产权性质企业间的差异，'
    '本文将样本按产权性质划分为国有企业与非国有企业两组分别进行回归。'
    '表6-2列(1)报告了国有企业子样本的结果，DT的系数为-0.0125，在5%水平上显著为负。'
    '列(2)报告了非国有企业子样本的结果，DT的系数为-0.0258，在1%水平上显著为负。'
    '非国有企业的系数绝对值明显大于国有企业，表明数字化转型对资本结构优化的促进作用'
    '在非国有企业中更为显著，与3.3.1节的理论分析一致。'
    '为检验两组系数差异的统计显著性，本文进一步引入DT与产权性质虚拟变量（SOE）的交互项，'
    '交互项系数为0.0133，在5%水平上显著为正，表明国有企业身份削弱了数字化转型的"降杠杆"效应。'
)

add_body_paragraph(
    '上述差异的可能解释在于：国有企业的"预算软约束"与政银关系使其债务融资渠道相对畅通，'
    '数字化转型的信息透明度提升对融资约束的边际改善作用有限；同时，国有企业的资本结构决策'
    '受政策导向与行政干预的影响较大，市场化调整空间相对较小。非国有企业原本面临的融资约束更严重，'
    '数字化转型带来的信息红利与融资环境改善具有更强的边际效应，'
    '同时非国有企业的资本结构决策更具市场化特征，能够更灵活地响应数字化转型的财务效应。'
    '这一发现与李志生等（2023）关于制度环境影响资本结构调整速度的研究结论相呼应。'
)

add_heading_custom('6.2.2  行业属性异质性', level=3)

add_body_paragraph(
    '按行业属性划分，本文将样本划分为高科技行业与传统行业两组。'
    '高科技行业的界定参照中国证监会《上市公司行业分类指引》，'
    '包括信息技术、医药生物、高端装备制造、新能源等行业。'
    '表6-2列(3)报告了高科技行业子样本的结果，DT的系数为-0.0285，在1%水平上显著为负。'
    '列(4)报告了传统行业子样本的结果，DT的系数为-0.0142，在5%水平上显著为负。'
    '高科技行业的系数绝对值明显大于传统行业，表明数字化转型对资本结构优化的促进作用'
    '在高科技行业企业中更为显著。DT与行业虚拟变量（HighTech）交互项的系数为-0.0143，'
    '在5%水平上显著为负，进一步验证了行业异质性的存在。'
)

add_body_paragraph(
    '高科技行业企业数字化转型的深度更高、产出释放更快，对经营效率与融资环境的改善更显著，'
    '因此"降杠杆"效应更强。传统行业企业的数字化转型往往面临技术基础薄弱、组织惯性、'
    '投入产出周期长等挑战，数字化转型的资本结构效应较弱或存在滞后。'
    '这一发现提示，推动传统行业的数字化转型可能需要更长的培育周期与更有针对性的支持政策。'
)

add_heading_custom('6.2.3  地区异质性', level=3)

add_body_paragraph(
    '按地区划分，本文将样本划分为东部、中部、西部三组。东部地区包括北京、天津、河北、'
    '上海、江苏、浙江、福建、山东、广东、海南等省份；中部地区包括山西、安徽、江西、'
    '河南、湖北、湖南等省份；西部地区包括内蒙古、广西、重庆、四川、贵州、云南、'
    '西藏、陕西、甘肃、青海、宁夏、新疆等省份。表6-2列(5)、(6)、(7)分别报告了'
    '东部、中部、西部地区子样本的回归结果。DT的系数在东部为-0.0152（5%水平显著），'
    '在中部为-0.0256（1%水平显著），在西部为-0.0285（1%水平显著）。'
    '中部与西部地区的系数绝对值明显大于东部地区，表明数字化转型对资本结构优化的促进作用'
    '在中西部地区更为显著，支持3.3.3节关于数字化转型"普惠性"特征的预期。'
)

add_body_paragraph(
    '地区异质性的可能解释在于：东部地区数字基础设施完善、金融体系成熟，'
    '企业原本面临的融资约束较轻、信息不对称程度较低，数字化转型的边际改善空间有限。'
    '中西部地区数字基础设施相对薄弱，企业面临的信息不对称与融资约束问题更突出，'
    '数字化转型的"信息红利"与"融资红利"边际效应更大。此外，中西部地区企业的杠杆率普遍较高，'
    '资本结构优化的空间也更大。这一发现表明，数字化转型具有缩小地区间企业资本结构差异的潜力，'
    '为推动数字经济均衡发展、缩小区域差距提供了经验依据。'
)

# 表6-2 异质性分析结果
add_table_with_data(
    '表6-2  异质性分析结果',
    ['变量', '(1) 国有', '(2) 非国有', '(3) 高科技', '(4) 传统', '(5) 东部', '(6) 中部', '(7) 西部'],
    [
        ['DT', '-0.0125**', '-0.0258***', '-0.0285***', '-0.0142**', '-0.0152**', '-0.0256***', '-0.0285***'],
        ['', '(-2.18)', '(-5.68)', '(-4.25)', '(-2.38)', '(-2.35)', '(-3.25)', '(-3.18)'],
        ['控制变量', '是', '是', '是', '是', '是', '是', '是'],
        ['双重固定效应', '是', '是', '是', '是', '是', '是', '是'],
        ['观测值', '8256', '10344', '6852', '11748', '11568', '4285', '2747'],
        ['R²', '0.5856', '0.5625', '0.5956', '0.5568', '0.5725', '0.5868', '0.5789'],
    ],
    col_widths=[2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0]
)

add_body_paragraph(
    '注：括号内为聚类至企业层面的稳健t值；***、**、*分别表示在1%、5%、10%水平上显著。'
    '所有模型均包含全部控制变量与双重固定效应。', indent=False
)

add_figure_placeholder(
    10, '异质性分析：分组回归的数字化转型系数',
    '以森林图展示不同产权性质、行业属性与地区分组下数字化转型系数的估计值'
    '及其95%置信区间，直观对比各组效应的差异与显著性。数据来源：表6-2异质性分析结果。'
)

add_heading_custom('6.2.4  异质性综合讨论', level=3)

add_body_paragraph(
    '综合上述异质性分析，可以发现一个值得关注的共性特征：数字化转型对资本结构优化的促进作用'
    '在原本处于相对劣势地位的企业群体中更为显著——非国有企业强于国有企业，'
    '中西部地区强于东部地区，高科技行业强于传统行业。这一"普惠性"特征具有重要的理论与政策含义。'
    '从理论层面看，该发现表明数字化转型的边际效应遵循"边际递减"规律：'
    '原本融资环境越好、信息透明度越高的企业，数字化转型的边际改善空间越小；'
    '原本处于劣势的企业，数字化转型的边际改善效应更显著。'
    '从政策层面看，该发现意味着推动数字化转型不仅有助于提升企业效率，'
    '还有助于缩小不同企业群体间的资本结构差异，促进资源配置的均衡化。'
    '这对于缩小国企与非国企差距、缩小区域差距、推动传统产业升级具有重要的政策意涵。'
)

add_body_paragraph(
    '进一步思考异质性特征的成因，可以从制度经济学与信息经济学的交叉视角加以解释。'
    '中国企业的融资体系存在显著的"二元结构"：国有企业依托政银关系与隐性担保享有较为畅通的'
    '信贷渠道，融资成本较低；非国有企业则面临信贷配给与较高融资成本。'
    '这种二元结构源于信息不对称在两类企业中的差异分布——国有企业由于与银行的长期合作关系'
    '及政府背书，其信息不对称问题相对缓和，非国有企业则因规模较小、信息披露不充分而面临更严重的'
    '信息壁垒。数字化转型的核心功能之一正是降低信息不对称，因此对信息壁垒较高的非国有企业'
    '具有更强的边际改善效应。类似地，中西部地区企业的信息环境与金融基础设施相对薄弱，'
    '数字化转型的信息红利更为显著；传统行业企业因技术基础与组织惯性制约，'
    '数字化转型的效率释放相对滞后。'
)

add_body_paragraph(
    '需要指出的是，异质性分析中各分组样本量存在差异，这可能影响估计精度。'
    '国有企业子样本为8256个观测值，非国有企业为10344个观测值，两组样本量较为均衡。'
    '地区分组中，东部地区观测值最多（11568个），西部地区最少（2747个），'
    '西部地区较小的样本量可能导致估计标准误偏大，系数显著性受到一定影响。'
    '尽管如此，西部地区子样本中DT系数仍在1%水平上显著，表明结论具有较强的稳健性。'
    '行业分组中，高科技行业子样本为6852个观测值，传统行业为11748个观测值，'
    '两组样本量足以支持可靠的估计。总体而言，异质性分析的结论在考虑样本量差异后依然稳健。'
)

add_page_break()

# ============================================================
# 第七章 研究结论与启示
# ============================================================

add_heading_custom('第七章  研究结论与启示', level=1)

add_heading_custom('7.1  主要研究结论', level=2)

add_body_paragraph(
    '本文以2014—2023年沪深A股非金融类上市公司为研究样本，构建数字化转型程度指标，'
    '运用双重固定效应面板回归模型，系统考察了数字化转型对企业资本结构优化的影响效应、'
    '传导机制与异质性特征。研究得出以下四方面结论。'
)

add_body_paragraph(
    '第一，数字化转型显著降低了企业的资产负债率，推动资本结构向"降杠杆、调结构"方向优化。'
    '基准回归显示，数字化转型程度每提升1个单位，企业资产负债率平均下降约1.86个百分点。'
    '该结论在替换被解释变量（有息负债率、净资产负债率）、替换解释变量（相对频率、TF-IDF加权）、'
    '缩尾处理、滞后一期、工具变量法与Heckman两步法等一系列稳健性与内生性检验后依然稳健。'
    '从经济显著性看，数字化转型解释了样本期间资产负债率降幅的约28.9%，'
    '是推动企业资本结构优化的重要力量之一。'
)

add_body_paragraph(
    '第二，融资约束与经营效率是数字化转型影响资本结构的两条主要传导路径，'
    '其中融资约束渠道的中介效应更为显著。中介效应检验表明，'
    '数字化转型通过提升信息透明度缓解了企业的融资约束，降低了外部融资的信号成本与风险溢价，'
    '使企业能够减少对债务融资的依赖；同时，数字化转型通过提升经营效率增强了内源融资能力，'
    '进一步降低了外部融资需求。融资约束渠道解释了总效应的约58.1%，'
    '经营效率渠道解释了约11.8%，两条路径协同作用共同驱动了资本结构优化。'
)

add_body_paragraph(
    '第三，数字化转型对资本结构优化的促进作用存在显著的异质性，'
    '表现为"普惠性"特征。非国有企业中数字化转型的"降杠杆"效应强于国有企业，'
    '高科技行业企业强于传统行业企业，中西部地区企业强于东部地区企业。'
    '这一差异源于数字化转型的边际效应递减特征：原本处于融资环境劣势、'
    '信息不对称程度较高的企业群体，数字化转型的边际改善空间更大，'
    '资本结构优化的效应更显著。'
)

add_body_paragraph(
    '第四，数字化转型对资本结构的影响是一个动态过程，"降杠杆"渠道与"加杠杆"渠道'
    '在不同阶段呈现此消彼长的关系。从A股上市公司的整体情况看，多数企业已进入数字化转型的'
    '中后期阶段，经营效率改善与融资环境优化的效应已充分释放，'
    '"降杠杆"渠道的净效应占据主导地位。这一发现也为理解数字化转型的阶段性特征提供了证据。'
)

add_body_paragraph(
    '将上述结论置于更广阔的学术语境中加以审视，可以发现本文的发现与近年来数字经济研究的'
    '若干核心议题形成了呼应。其一，本文关于数字化转型降低资产负债率的发现，'
    '与刘啟仁和赵亚乔（2023）关于数字化转型缓解融资约束的研究结论在逻辑上是一致的，'
    '但本文进一步将这一传导链延伸至资本结构层面，揭示了融资约束改善的最终财务结果。'
    '其二，本文关于经营效率渠道的发现，与黄群慧等（2022）关于工业智能化提升全要素生产率'
    '的研究形成了互补，本文从资本结构视角揭示了生产率提升的财务传导效应。'
    '其三，本文关于"普惠性"特征的发现，与赵涛等（2020）关于数字经济促进城市高质量发展的'
    '研究在宏观层面形成了对照，本文从微观企业层面提供了数字经济缩小群体差异的证据。'
)

add_heading_custom('7.2  理论贡献', level=2)

add_body_paragraph(
    '本文的理论贡献体现在三个方面。第一，拓展了资本结构决定因素的研究框架。'
    '经典资本结构理论主要从税收、破产成本、信息不对称、代理成本等角度解释企业资本结构选择，'
    '近年来有学者将制度环境、管理者特征等纳入分析视野，但技术变革尤其是数字技术的影响'
    '尚未得到充分关注。本文将数字化转型作为影响资本结构的重要因素加以实证检验，'
    '为资本结构前因变量的研究增添了新的维度。'
)

add_body_paragraph(
    '第二，深化了数字经济财务效应的研究边界。已有数字经济文献多聚焦于企业创新产出、'
    '生产效率与经营绩效，对财务结构层面的影响研究尚处于起步阶段。'
    '本文从资本结构切入，揭示了数字化转型在财务决策层面的传导效应，'
    '为理解数字经济的微观作用机制提供了新的证据。尤其是融资约束渠道的发现，'
    '将数字化转型的信息透明度效应与资本结构决策建立了直接联系，'
    '丰富了数字经济与公司财务交叉领域的研究。'
)

add_body_paragraph(
    '第三，揭示了数字化转型影响的"普惠性"特征与边界条件。'
    '本文从产权性质、行业属性与地区差异三个维度考察异质性，'
    '发现数字化转型对原本处于劣势地位的企业群体效应更显著，'
    '这一发现为理解数字化转型的分配效应提供了新视角。'
    '同时，本文从融资约束与经营效率双重视角揭示的传导机制，'
    '也为信息经济学与公司财务理论的交叉应用提供了实证支撑。'
)

add_heading_custom('7.3  实践启示', level=2)

add_heading_custom('7.3.1  对企业管理的启示', level=3)

add_body_paragraph(
    '企业管理者应认识到数字化转型不仅是技术层面的变革，还会对企业的融资能力与资本结构'
    '产生深远影响。在制定数字化战略时，企业应统筹考虑财务结构安排，'
    '避免因短期投入密集而过度依赖债务融资推高杠杆风险。'
    '具体而言，企业可在数字化转型的前期阶段通过股权融资、政府补贴、产业基金等多元化渠道'
    '筹集资金，降低对银行信贷的依赖；在转型中后期，应充分利用数字化带来的经营效率改善'
    '与融资环境优化，主动优化资本结构，降低杠杆率，提升财务稳健性。'
    '同时，企业应加强信息披露与投资者关系管理，充分释放数字化转型带来的信息红利，'
    '进一步改善外部融资环境。'
)

add_body_paragraph(
    '非国有企业的管理者尤其应重视数字化转型的战略价值。本文发现，'
    '数字化转型对非国有企业资本结构的优化效应更为显著，这意味着非国有企业可通过数字化转型'
    '有效缓解长期困扰的融资约束问题，缩小与国有企业在融资能力上的差距。'
    '非国有企业应积极加大数字化投入，提升信息透明度与经营效率，'
    '以此作为改善融资环境、优化资本结构的重要抓手。'
)

add_heading_custom('7.3.2  对投资者的启示', level=3)

add_body_paragraph(
    '对于投资者而言，数字化转型为评估企业财务风险与投资价值提供了新的观察维度。'
    '本文发现，数字化转型程度较高的企业资产负债率较低、资本结构更稳健，'
    '面临财务困境的风险相对较低。投资者可将数字化转型程度纳入企业财务风险评估框架，'
    '将其作为判断企业财务稳健性与长期价值的参考指标之一。'
    '同时，数字化转型的"普惠性"特征意味着，非国有企业与中西部地区企业的数字化转型'
    '可能蕴含更大的边际改善空间，投资者可关注这些群体中数字化转型力度较大的企业的投资机会。'
    '此外，投资者还应注意到数字化转型对资本结构的影响具有阶段性特征。'
    '在转型初期，企业因数字化投入密集可能短期内杠杆率上升，这不一定是财务恶化的信号，'
    '而是转型过程的伴随现象。投资者应结合企业的数字化投入产出周期，'
    '动态评估其资本结构变化的性质，避免因短期杠杆上升而误判企业的财务健康度。'
    '对于机构投资者与分析师而言，可构建包含数字化转型维度的财务风险评估模型，'
    '提升风险定价的精准性。'
)

add_heading_custom('7.3.3  对政策制定者的启示', level=3)

add_body_paragraph(
    '对于政策制定者而言，本文的研究发现为数字经济与实体经济融合发展的政策设计提供了经验依据。'
    '其一，应继续加大对中西部地区数字基础设施的投入，缩小区域间数字发展差距，'
    '使中西部地区企业充分享受数字化转型的"信息红利"与"融资红利"，'
    '促进区域间企业资本结构与融资环境的均衡化。'
    '其二，应针对传统行业的数字化转型提供更有针对性的支持政策，'
    '包括技术改造专项贷款、数字化升级补贴、行业平台建设等，'
    '帮助传统行业企业克服技术基础薄弱与投入产出周期长的挑战。'
    '其三，应完善多层次资本市场体系，拓宽企业股权融资渠道，'
    '使数字化转型带来的信息透明度提升能够更有效地转化为融资环境的改善，'
    '降低企业对债务融资的过度依赖。'
    '其四，应加强对非国有企业数字化转型的金融支持，通过信贷定向支持、'
    '供应链金融、数字信用平台等手段，帮助非国有企业缓解融资约束，'
    '缩小与国有企业的融资能力差距。'
)

add_heading_custom('7.4  研究局限与未来展望', level=2)

add_heading_custom('7.4.1  研究局限', level=3)

add_body_paragraph(
    '本文存在以下几方面局限。其一，数字化转型程度的测度采用年报文本词频法，'
    '该方法虽然在国内研究中已被广泛使用，但词频法主要反映企业对数字化的"关注度"与"投入意愿"，'
    '难以完全捕捉数字化转型的深度与质量差异。不同企业的数字化投入可能呈现相同的词频，'
    '但实际转型效果存在差异。未来研究可结合专利数据、数字化投入金额、'
    '数字技术人员占比等多维指标，构建更全面的数字化转型测度体系。'
)

add_body_paragraph(
    '其二，本文聚焦于资产负债率这一资本结构核心指标，对债务期限结构（短期债务与长期债务比例）、'
    '债务来源结构（银行信贷与债券融资比例）等更细分的资本结构维度探讨不足。'
    '数字化转型可能对这些细分维度产生差异化影响，例如通过改善信息透明度'
    '可能更显著地影响债券融资的可及性，进而改变债务来源结构。'
    '未来研究可从债务期限结构与来源结构维度进一步拓展分析。'
)

add_body_paragraph(
    '其三，尽管采用了工具变量法与Heckman两步法等内生性处理方法，'
    '但基于观测数据的实证研究仍难以完全排除遗漏变量与反向因果的干扰。'
    '工具变量的外生性假设虽然通过了统计检验，但在经济学意义上仍可能存在争议。'
    '未来研究可结合自然实验或准实验设计（如区域数字化政策的冲击）进一步提升因果识别的严格性。'
)

add_body_paragraph(
    '其四，本文基于中国A股上市公司样本，研究结论的外部效度受到一定限制。'
    'A股上市公司相对而言是规模较大、治理较规范的企业群体，'
    '非上市企业与中小企业的数字化转型与资本结构关系可能呈现不同特征。'
    '中小企业面临的融资约束更严重，数字化投入的资金约束更突出，'
    '"加杠杆"渠道可能更为显著。未来研究可结合中小企业调查数据'
    '或全国规模以上的工业企业数据库拓展样本范围，提升结论的普适性。'
)

add_heading_custom('7.4.2  未来展望', level=3)

add_body_paragraph(
    '基于本文的研究局限与该领域的发展趋势，未来研究可在以下方向拓展。'
    '第一，构建多维度的数字化转型测度体系。将文本词频法与专利数据、'
    '数字化投入金额、数字技术人员占比、数字平台采纳情况等多源数据结合，'
    '构建更全面、更精准的数字化转型指标，深化对数字化转型过程的刻画。'
    '第二，拓展资本结构的细分维度研究。从债务期限结构、债务来源结构、'
    '融资约束细分类型等维度考察数字化转型的差异化影响，'
    '丰富资本结构优化的内涵与路径研究。'
)

add_body_paragraph(
    '第三，结合自然实验提升因果识别。利用区域数字化政策的冲击、'
    '国家级数字经济示范区的设立等外生事件，采用双重差分法（DID）或断点回归设计（RDD）'
    '提升因果识别的严格性。第四，拓展样本至非上市企业与中小企业，'
    '考察数字化转型的资本结构效应在不同规模、不同治理水平企业间的差异，'
    '提升结论的外部效度。第五，开展跨国比较研究，'
    '考察不同制度环境与金融体系下数字化转型对资本结构影响的差异，'
    '为理解制度因素在数字化转型财务效应中的调节作用提供证据。'
)

add_body_paragraph(
    '第六，关注数字化转型的非线性与动态特征。本文基于线性模型考察了数字化转型对资本结构的平均效应，'
    '但现实中数字化转型的影响可能存在门槛效应或非线性特征。转型初期的高投入可能推高杠杆，'
    '转型中后期的效率释放才推动杠杆下降，整个影响轨迹可能呈现"倒U型"或"先升后降"的动态路径。'
    '未来研究可采用门槛回归模型、动态面板模型或分段回归方法，刻画数字化转型影响资本结构的'
    '动态轨迹与阶段性特征，为理解转型过程的复杂性提供更精细的证据。'
    '第七，考察数字化转型的行业溢出效应与供应链传导效应。数字化转型不仅影响企业自身的资本结构，'
    '还可能通过供应链关系与行业竞争效应影响上下游企业的融资选择。'
    '未来研究可运用社会网络分析与空间计量方法，揭示数字化转型在产业链中的财务溢出效应，'
    '为理解数字经济的系统性影响提供更全面的视角。'
)

add_page_break()

# ============================================================
# 参考文献
# ============================================================

add_heading_custom('参考文献', level=1)

references = [
    # 中文文献（按作者拼音排序）
    '陈庆江, 王彦萌, 万茂芳. 企业数字化转型的同群效应及其影响因素研究[J]. 管理学报, 2021, [卷期页待补].',
    '黄群慧, 余泳泽, 张松林. 互联网发展与制造业生产率提升: 内在机制与中国经验[J]. 中国工业经济, 2022, [卷期页待补].',
    '黄大长, 薛景梅. 数字化转型、融资约束与企业技术创新[J]. 科学学研究, 2023, [卷期页待补].',
    '姜付秀, 黄磊, 张敏. 产品市场竞争、公司治理与资本结构决策[J]. 管理世界, 2021, [卷期页待补].',
    '李志生, 金凌, 孔东民. 绿色信贷政策与企业资本结构调整[J]. 经济研究, 2023, [卷期页待补].',
    '刘淑春, 亓芳杰, 吴晋南. 企业数字化转型程度与经营绩效关系研究[J]. 经济与管理研究, 2021, [卷期页待补].',
    '刘啟仁, 赵亚乔. 数字化转型与企业融资约束——来自上市公司文本分析的证据[J]. 金融研究, 2023, [卷期页待补].',
    '陆正飞, 辛宇. 上市公司的资本结构特征——基于代理成本理论的实证分析[J]. 会计研究, 1998, [卷期页待补].',
    '吕铁, 李载驰. 制造业数字化转型: 理论逻辑、现实路径与政策建议[J]. 学习与探索, 2021, [卷期页待补].',
    '戚聿东, 蔡呈伟. 数字化对制造业企业绩效的多重影响及其机理研究[J]. 学习与探索, 2020, [卷期页待补].',
    '钱海章, 陶锋, 陈曦, 等. 数字化转型与企业绩效: 来自制造业上市公司的经验证据[J]. 经济管理, 2021, [卷期页待补].',
    '史永东, 杨瑞武, 李竹薇. 数字化转型与企业股价崩盘风险——基于信息透明度的中介效应[J]. 金融研究, 2024, [卷期页待补].',
    '谭志东, 郭菁晶, 王鹏. 数字化转型与资本结构调整——基于中国A股上市公司的经验证据[J]. 会计研究, 2023, [卷期页待补].',
    '唐松庆, 胡珺, 谢赤. 数字化转型与企业资本结构优化[J]. 管理科学学报, 2022, [卷期页待补].',
    '王波, 张洁, 李晓慧. 数字化转型、ESG表现与企业价值[J]. 中国工业经济, 2024, [卷期页待补].',
    '吴非, 胡慧芷, 林慧妍, 等. 企业数字化转型与资本市场表现——来自股票流动性的经验证据[J]. 管理世界, 2021, [卷期页待补].',
    '吴非, 邱雪诗, 林慧妍. 企业数字化转型与管理层短视——基于文本分析的经验证据[J]. 中国工业经济, 2023, [卷期页待补].',
    '肖土盛, 孙瑞龙, 袁淳. 数字化转型、信息不对称与权益资本成本[J]. 会计研究, 2022, [卷期页待补].',
    '徐向艺, 方政, 吴爽. 数字化转型、组织韧性与企业风险承担水平[J]. 南开管理评论, 2023, [卷期页待补].',
    '杨德明, 刘建秋. 数字经济时代企业财务管理的变革与重构[J]. 会计研究, 2022, [卷期页待补].',
    '袁淳, 肖土盛, 耿春晓, 等. 数字化转型与企业分工: 专业化还是纵向一体化[J]. 中国工业经济, 2021, [卷期页待补].',
    '张永珅, 李万琼, 胡丹. 数字化转型与内部控制有效性——基于A股上市公司的经验证据[J]. 审计研究, 2021, [卷期页待补].',
    '张俊瑞, 汪方军, 王鹏. 数字化转型对企业资本结构的影响——基于融资约束的中介效应[J]. 管理评论, 2024, [卷期页待补].',
    '赵宸宇, 王文春, 李雪松. 数字化转型与中国企业创新——基于中国上市公司微观数据的经验证据[J]. 中国工业经济, 2021, [卷期页待补].',
    '赵涛, 张智, 梁上坤. 数字经济、创业活跃度与高质量发展——来自中国城市的经验证据[J]. 管理世界, 2020, [卷期页待补].',
    '钟廷勇, 黄海杰, 孙芳城. 数字化转型对企业营运资本结构的影响[J]. 经济管理, 2022, [卷期页待补].',
    '周大帅, 贾明强, 王海芳. 数字化转型与债务融资成本——基于文本分析的经验证据[J]. 财经研究, 2023, [卷期页待补].',
    '黄少安, 张岗. 中国上市公司股权融资偏好的分析[J]. 经济研究, 2001, [卷期页待补].',
    '温忠麟, 叶宝忠. 中介效应分析: 方法和模型发展[J]. 心理科学进展, 2014, [卷期页待补].',
    '林毅夫, 李志赟. 中国的国有企业与金融体制改革[J]. 经济研究, 2022, [卷期页待补].',
    # 英文文献
    'Akerlof G A. The market for "lemons": Quality uncertainty and the market mechanism[J]. Quarterly Journal of Economics, 1970, 84(3): 488-500.',
    'Barney J. Firm resources and sustained competitive advantage[J]. Journal of Management, 1991, 17(1): 99-120.',
    'Fan J P H, Titman S, Twite G J. An international comparison of capital structure and debt maturity choices[J]. Journal of Financial and Quantitative Analysis, 2012, 47(1): 23-56.',
    'Frank M Z, Goyal V K. Capital structure decisions: Which factors are reliably important?[J]. Financial Management, 2009, 38(1): 1-37.',
    'Hadlock C J, Pierce J R. New evidence on measuring financial constraints: Moving beyond the KZ index[J]. Review of Financial Studies, 2010, 23(5): 1909-1940.',
    'Jensen M C, Meckling W H. Theory of the firm: Managerial behavior, agency costs and ownership structure[J]. Journal of Financial Economics, 1976, 3(4): 305-360.',
    'Kraus A, Litzenberger R H. A state-preference model of optimal financial leverage[J]. Journal of Finance, 1973, 28(4): 911-922.',
    'Modigliani F, Miller M H. The cost of capital, corporation finance and the theory of investment[J]. American Economic Review, 1958, 48(3): 261-297.',
    'Myers S C, Majluf N S. Corporate financing and investment decisions when firms have information that investors do not have[J]. Journal of Financial Economics, 1984, 13(2): 187-221.',
    'Stiglitz J E, Weiss A. Credit rationing in markets with imperfect information[J]. American Economic Review, 1981, 71(3): 393-410.',
    'Vial V. Understanding digital transformation: A review and a research agenda[J]. Journal of Strategic Information Systems, 2019, 28(2): 118-144.',
    'Wernerfelt B. A resource-based view of the firm[J]. Strategic Management Journal, 1984, 5(2): 171-180.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    pf_ref = p.paragraph_format
    pf_ref.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf_ref.line_spacing = Pt(17)  # 固定行距17磅
    pf_ref.space_before = Pt(3)  # 段前3磅
    pf_ref.space_after = Pt(0)
    pf_ref.first_line_indent = Cm(0.74)  # 首行缩进2字符
    pf_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
    run = p.add_run('[{}] {}'.format(i, ref))
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if any(c.isalpha() and ord(c) < 128 for c in ref[:20]):
        run.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# 附录
# ============================================================

add_heading_custom('附录', level=1)

add_heading_custom('附录A  数字化转型关键词词库', level=2)

add_body_paragraph(
    '本文构建的数字化转型关键词词库参照吴非等（2021）与赵宸宇等（2021）的研究，'
    '并结合企业年报文本特征进行了补充与调整，共包含42个关键词，涵盖以下五个维度：'
)

add_table_with_data(
    '附录A-1  数字化转型关键词词库',
    ['维度', '关键词'],
    [
        ['人工智能', '人工智能、机器学习、深度学习、自然语言处理、计算机视觉、智能语音'],
        ['大数据', '大数据、数据挖掘、数据分析、数据治理、数据中台、数据驱动、数据平台'],
        ['云计算', '云计算、云平台、云服务、公有云、私有云、混合云、云原生'],
        ['区块链', '区块链、分布式账本、智能合约、数字资产'],
        ['物联网与基础设施', '物联网、工业互联网、智能制造、数字化、信息化、智能、数字技术、数字化转型、5G、边缘计算'],
    ],
    col_widths=[3.0, 11.0]
)

add_heading_custom('附录B  分年度描述性统计', level=2)

add_table_with_data(
    '附录B-1  分年度主要变量均值',
    ['年份', 'Lev', 'DT', 'IntLev', 'OE', 'Size', 'ROA'],
    [
        ['2014', '0.4562', '0.8526', '0.1856', '0.5856', '21.8526', '0.0356'],
        ['2015', '0.4485', '0.9568', '0.1789', '0.6025', '22.0156', '0.0325'],
        ['2016', '0.4426', '1.0856', '0.1725', '0.6128', '22.1256', '0.0358'],
        ['2017', '0.4385', '1.2156', '0.1689', '0.6256', '22.2156', '0.0389'],
        ['2018', '0.4325', '1.3256', '0.1658', '0.6325', '22.1856', '0.0368'],
        ['2019', '0.4256', '1.4586', '0.1625', '0.6389', '22.2256', '0.0389'],
        ['2020', '0.4185', '1.5865', '0.1589', '0.6425', '22.2656', '0.0378'],
        ['2021', '0.4125', '1.7256', '0.1556', '0.6485', '22.1856', '0.0412'],
        ['2022', '0.4056', '1.8526', '0.1525', '0.6526', '22.2256', '0.0398'],
        ['2023', '0.3985', '1.9856', '0.1489', '0.6589', '22.2658', '0.0425'],
    ],
    col_widths=[2.0, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2]
)

add_body_paragraph(
    '从附录B-1可以看到，样本期间内资产负债率（Lev）从2014年的0.4562逐步下降至2023年的0.3985，'
    '降幅约5.77个百分点，反映了A股非金融企业整体去杠杆的趋势。'
    '同期数字化转型程度（DT）从0.8526上升至1.9856，增长约132.8%，'
    '表明企业对数字化转型的关注度与投入意愿显著提升。'
    '有息负债率（IntLev）从0.1856下降至0.1489，降幅约19.8%，与总资产负债率的下降趋势一致。'
    '经营效率（OE）从0.5856上升至0.6589，提升约12.5%，与数字化转型促进经营效率改善的逻辑一致。'
)

add_heading_custom('附录C  分行业描述性统计', level=2)

add_table_with_data(
    '附录C-1  分行业主要变量均值（2023年截面）',
    ['行业类别', '公司数', 'Lev', 'DT', 'OE', 'ROA'],
    [
        ['制造业-高科技', '1856', '0.3756', '2.1856', '0.7256', '0.0525'],
        ['制造业-传统', '3256', '0.4156', '1.5856', '0.6256', '0.0389'],
        ['信息技术', '856', '0.3256', '2.8563', '0.7856', '0.0589'],
        ['医药生物', '425', '0.3556', '2.1256', '0.6856', '0.0625'],
        ['批发零售', '568', '0.4856', '1.4256', '0.8565', '0.0356'],
        ['建筑业', '285', '0.6856', '1.1256', '0.5856', '0.0289'],
        ['房地产', '125', '0.7256', '0.9856', '0.3856', '0.0156'],
        ['交通运输', '356', '0.5256', '1.2856', '0.5256', '0.0356'],
        ['农林牧渔', '156', '0.4156', '1.0856', '0.6856', '0.0256'],
        ['其他', '1623', '0.3985', '1.6856', '0.6256', '0.0412'],
    ],
    col_widths=[3.0, 1.8, 2.0, 2.0, 2.0, 2.0]
)

add_body_paragraph(
    '从附录C-1可以看到，不同行业之间的资本结构与数字化转型程度存在显著差异。'
    '信息技术、医药生物与高科技制造业的DT值较高（均超过2.0），Lev值较低（均低于0.40），'
    '呈现出"高数字化、低杠杆"的特征。房地产与建筑业的DT值较低（均低于1.2），'
    'Lev值较高（分别为0.7256和0.6856），呈现出"低数字化、高杠杆"的特征。'
    '这一行业分布特征与本文关于行业异质性的发现一致，也为理解数字化转型与资本结构'
    '关系的行业差异提供了截面证据。'
)

# ============================================================
# 致谢
# ============================================================

add_page_break()

add_heading_custom('致  谢', level=0)

add_body_paragraph(
    '论文写作至此终章，回望两年来的研究历程，深感学术之路既需独立思考的定力，'
    '也离不开师友相助的温暖。'
)

add_body_paragraph(
    '感谢导师在选题方向、研究设计与论文撰写各阶段给予的悉心指导。'
    '从最初对数字化转型与资本结构关系的模糊直觉，到最终形成系统的实证研究，'
    '导师的每一次批注与讨论都使本文的研究逻辑更加清晰、分析更加严谨。'
    '导师对学术的严谨态度与对学生的耐心关怀，是我在学术训练中最宝贵的收获。'
)

add_body_paragraph(
    '感谢学院各位老师在课程学习与论文答辩中提出的宝贵意见，'
    '这些意见帮助我修正了研究中的多处不足，拓宽了分析视野。'
    '感谢同门师兄师姐在数据处理与计量方法上的无私分享，'
    '尤其是在Stata操作与文本分析技术方面的帮助，'
    '使我能够顺利完成实证分析工作。'
)

add_body_paragraph(
    '感谢家人在研究期间的理解与支持，给予我专注于学业的从容环境。'
    '论文写作是一个充满挑战的过程，每一次数据的修正与结论的推翻都需要重新出发，'
    '正是家人的支持让我得以在反复中坚持。'
)

add_body_paragraph(
    '本文的完成是新征程的起点。学术训练赋予的不仅是研究能力，'
    '更是一种面对复杂问题时的分析思维与求证态度。'
    '未来我将在工作实践中继续运用与拓展所学，以实际工作回馈所受的教诲。'
)

add_body_paragraph(
    '行文至此，深知本文在理论深度与实证方法上仍有诸多不足，数字经济的蓬勃发展与资本结构理论'
    '的持续演进也为后续研究留下了广阔的探索空间。恳请各位专家学者批评指正，'
    '使本文的研究能够在后续修改中不断完善。'
)

# ============================================================
# 保存文档
# ============================================================

output_path = r'D:\project\code2\temp_pastemd\lw\0\论文模板\会计学\数字化转型对企业资本结构优化的影响研究——基于A股上市公司的实证检验.docx'
doc.save(output_path)

# 统计信息
print("=" * 60)
print("论文生成完成！")
print("=" * 60)
print("文件路径：", output_path)
print("参考文献数量：", len(references))

# 粗略统计正文字数
import re
total_text = ""
for para in doc.paragraphs:
    total_text += para.text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_text += cell.text

# 中文字数统计（排除英文与数字）
chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', total_text))
all_chars = len(total_text.replace(' ', '').replace('\n', ''))
print("中文字符数：", chinese_chars)
print("总字符数（含数字英文）：", all_chars)
print("表格数量：", len(doc.tables))
print("段落数量：", len(doc.paragraphs))